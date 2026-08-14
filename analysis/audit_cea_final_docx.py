#!/usr/bin/env python3
"""Structural audit for the revised CEA manuscript and supplement DOCX files."""

from __future__ import annotations

import json
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION = ROOT / "manuscript/cea_v29_submission"
QA_MAIN = ROOT / "results/cea_final_revision/docx_qa3/manuscript"
QA_SUPP = ROOT / "results/cea_final_revision/docx_qa4/supplement"
OUT = ROOT / "results/cea_final_revision"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def xml_root(path: Path) -> etree._Element:
    with ZipFile(path) as archive:
        return etree.fromstring(archive.read("word/document.xml"))


def inspect(label: str, docx: Path, pdf: Path, expected_tables: int,
            expected_images: int, expected_pages: int) -> dict[str, object]:
    doc = Document(docx)
    root = xml_root(docx)
    section = doc.sections[0]
    text = "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    widths = [int(value) for value in root.xpath(".//w:tblW/@w:w", namespaces=NS)]
    indents = [int(value) for value in root.xpath(".//w:tblInd/@w:w", namespaces=NS)]
    exact_heights = root.xpath(".//w:trHeight[@w:hRule='exact']", namespaces=NS)
    pages = len(PdfReader(pdf).pages)
    checks = {
        "docx_exists": docx.is_file() and docx.stat().st_size > 0,
        "pdf_exists": pdf.is_file() and pdf.stat().st_size > 0,
        "letter_page": section.page_width.twips == 12240 and section.page_height.twips == 15840,
        "one_inch_margins": all(
            value.twips == 1440 for value in
            [section.top_margin, section.bottom_margin, section.left_margin, section.right_margin]
        ),
        "table_count": len(doc.tables) == expected_tables,
        "image_count": len(doc.inline_shapes) == expected_images,
        "pdf_page_count": pages == expected_pages,
        "table_widths": not widths or all(value == 9360 for value in widths),
        "table_indents": not indents or all(value == 120 for value in indents),
        "no_exact_row_heights": len(exact_heights) == 0,
        "no_revision_history": not any(
            token in text.lower() for token in
            ["submission audit draft", "superseded v20", "obsolete v20", "manuscript version:"]
        ),
        "no_unresolved_markers": not any(
            token in text for token in ["[[", "]]", "{{", "}}", "[AUTHOR ACTION"]
        ),
    }
    return {
        "label": label,
        "status": "PASS" if all(checks.values()) else "FAIL",
        "docx": str(docx),
        "pdf": str(pdf),
        "pages": pages,
        "tables": len(doc.tables),
        "images": len(doc.inline_shapes),
        "checks": checks,
    }


def main() -> int:
    OUT.mkdir(parents=True, exist_ok=True)
    main_result = inspect(
        "manuscript",
        SUBMISSION / "PlumSPECTRA_CEA_manuscript_revised.docx",
        QA_MAIN / "PlumSPECTRA_CEA_manuscript_revised.pdf",
        expected_tables=0,
        expected_images=6,
        expected_pages=23,
    )
    supplement_result = inspect(
        "supplement",
        SUBMISSION / "PlumSPECTRA_CEA_supplement_revised.docx",
        QA_SUPP / "PlumSPECTRA_CEA_supplement_revised.pdf",
        expected_tables=7,
        expected_images=25,
        expected_pages=33,
    )
    (OUT / "docx_main_layout_final.json").write_text(
        json.dumps(main_result, indent=2), encoding="utf-8"
    )
    (OUT / "docx_supplement_layout_final.json").write_text(
        json.dumps(supplement_result, indent=2), encoding="utf-8"
    )
    status = main_result["status"] == supplement_result["status"] == "PASS"
    print(json.dumps({"status": "PASS" if status else "FAIL",
                      "main": main_result, "supplement": supplement_result}, indent=2))
    return 0 if status else 1


if __name__ == "__main__":
    sys.exit(main())
