from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

import pandas as pd
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TARGET_ORDER = [
    "skin_break_force_g_mean",
    "skin_break_displacement_raw_mean",
    "skin_break_drop_g_mean",
    "flesh_force_mean_g_mean",
    "force_at_6_rawpos_g_mean",
    "loading_stiffness_g_per_rawpos_mean",
    "loading_work_g_rawpos_mean",
    "post_break_work_g_rawpos_mean",
    "adhesive_force_g_mean",
]
LABELS_ZH = {
    "skin_break_force_g_mean": "表皮破裂力",
    "skin_break_displacement_raw_mean": "破裂位移",
    "skin_break_drop_g_mean": "破裂后力降",
    "flesh_force_mean_g_mean": "果肉平均阻力",
    "force_at_6_rawpos_g_mean": "原始位移6处力",
    "loading_stiffness_g_per_rawpos_mean": "加载刚度",
    "loading_work_g_rawpos_mean": "加载功",
    "post_break_work_g_rawpos_mean": "破裂后功",
    "adhesive_force_g_mean": "黏附力",
}
UNITS = {
    "skin_break_force_g_mean": "g-force",
    "skin_break_displacement_raw_mean": "raw_position_unit",
    "skin_break_drop_g_mean": "g-force",
    "flesh_force_mean_g_mean": "g-force",
    "force_at_6_rawpos_g_mean": "g-force",
    "loading_stiffness_g_per_rawpos_mean": "g/raw_position_unit",
    "loading_work_g_rawpos_mean": "g*raw_position_unit",
    "post_break_work_g_rawpos_mean": "g*raw_position_unit",
    "adhesive_force_g_mean": "g-force",
}
ICC_ENDPOINT = {
    "skin_break_force_g_mean": "skin_break_force_g",
    "skin_break_displacement_raw_mean": "skin_break_displacement_raw",
    "skin_break_drop_g_mean": "skin_break_drop_g",
    "flesh_force_mean_g_mean": "flesh_force_mean_g",
    "force_at_6_rawpos_g_mean": "force_at_6_rawpos_g",
    "loading_stiffness_g_per_rawpos_mean": "loading_stiffness_g_per_rawpos",
    "loading_work_g_rawpos_mean": "loading_work_g_rawpos",
    "post_break_work_g_rawpos_mean": "post_break_work_g_rawpos",
    "adhesive_force_g_mean": "adhesive_force_g",
}


def fmt(value: float, digits: int = 3) -> str:
    return f"{float(value):.{digits}f}"


def md_table(headers: list[str], rows: list[list[object]]) -> str:
    clean = lambda value: str(value).replace("|", "\\|").replace("\n", " ")
    output = ["| " + " | ".join(map(clean, headers)) + " |", "|" + "|".join(["---"] * len(headers)) + "|"]
    output.extend("| " + " | ".join(clean(value) for value in row) + " |" for row in rows)
    return "\n".join(output)


def dataframe_lookup(frame: pd.DataFrame, target: str, **filters: object) -> pd.Series:
    mask = frame["target"].eq(target)
    for column, value in filters.items():
        mask &= frame[column].eq(value)
    selected = frame.loc[mask]
    if len(selected) != 1:
        raise ValueError(f"Expected one row for {target}/{filters}, found {len(selected)}")
    return selected.iloc[0]


def build_markdown(project: Path) -> str:
    release = project.parent / "NIRs_plums_research_ready_en_v1.2.0"
    review = project / "results" / "project_review"

    cohort_counts = pd.read_csv(review / "cultivar_cohort_counts.csv")
    validity = pd.read_csv(review / "endpoint_transfer_validity_metrics.csv")
    model = pd.read_csv(project / "results" / "texture_prediction_analysis" / "texture_prediction_model_metrics.csv")
    random_metrics = pd.read_csv(project / "results" / "models" / "texture_pls_random_analysis" / "repeat_metrics.csv")
    random_mean = random_metrics.groupby("target", observed=True)["r2"].mean()
    fewshot = pd.read_csv(project / "results" / "models" / "texture_fewshot_analysis" / "fewshot_summary.csv")
    axis = pd.read_csv(project / "results" / "models" / "texture_axes_analysis" / "texture_axis_pooled_metrics.csv")
    axis_pls_fewshot = pd.read_csv(project / "results" / "models" / "texture_axis_fewshot_analysis" / "fewshot_summary.csv")
    axis_cnn_fewshot = pd.read_csv(project / "results" / "models" / "texture_axis_fewshot_cnn_analysis" / "fewshot_summary.csv")
    reliability = pd.read_csv(project / "results" / "texture_qc" / "tables" / "texture_qc_reliability_before_after.csv")
    sensitivity = pd.read_csv(review / "cohort_sensitivity_metrics.csv")
    audit = json.loads((release / "quality_control" / "independent_audit_report.json").read_text(encoding="utf-8"))
    qc_summary = json.loads((project / "results" / "texture_qc" / "texture_qc_impact_summary.json").read_text(encoding="utf-8"))
    critical = json.loads((review / "critical_validity_summary.json").read_text(encoding="utf-8"))

    cultivar_rows = []
    for row in cohort_counts.itertuples(index=False):
        cultivar_rows.append(
            [
                row.cultivar_ascii,
                row.cultivar_original,
                int(row.matched),
                int(row.analysis),
                int(row.strict),
                int(row.strict_complete),
            ]
        )
    cultivar_rows.append(
        [
            "合计",
            "-",
            int(cohort_counts["matched"].sum()),
            int(cohort_counts["analysis"].sum()),
            int(cohort_counts["strict"].sum()),
            int(cohort_counts["strict_complete"].sum()),
        ]
    )
    cultivar_table = md_table(
        ["English label", "原始标签", "匹配果", "主分析", "严格质构", "严格完整发布"], cultivar_rows
    )

    endpoint_rows = []
    fewshot_rows = []
    endpoint_definition_rows = []
    for target in TARGET_ORDER:
        val = validity.loc[validity["target"].eq(target)].iloc[0]
        pls = dataframe_lookup(model, target, model="PLSR")
        cnn = dataframe_lookup(model, target, model="1D-CNN")
        icc = reliability.loc[
            reliability["endpoint"].eq(ICC_ENDPOINT[target])
            & reliability["cohort"].eq("High-confidence analysis cohort")
        ].iloc[0]
        f0 = dataframe_lookup(fewshot, target, shots=0)
        f5 = dataframe_lookup(fewshot, target, shots=5)
        f20 = dataframe_lookup(fewshot, target, shots=20)
        endpoint_rows.append(
            [
                LABELS_ZH[target],
                fmt(icc.icc_a1),
                fmt(random_mean[target]),
                fmt(pls.r2),
                fmt(cnn.r2),
                fmt(val.within_cultivar_centered_r2),
                fmt(val.median_within_cultivar_r),
            ]
        )
        fewshot_rows.append(
            [
                LABELS_ZH[target],
                fmt(f0.r2_mean),
                fmt(f5.r2_mean),
                fmt(f5.ccc_mean),
                fmt(f20.r2_mean),
                fmt(f20.ccc_mean),
            ]
        )
        endpoint_definition_rows.append(
            [LABELS_ZH[target], target, UNITS[target], "两次穿刺特征均值；直接NIR预测目标"]
        )
    endpoint_table = md_table(
        ["端点", "ICC", "随机R2", "LOCO PLSR R2", "LOCO CNN R2", "品种内中心化R2", "品种内r中位数"],
        endpoint_rows,
    )
    fewshot_table = md_table(
        ["端点", "0-shot R2", "5-shot R2", "5-shot CCC", "20-shot R2", "20-shot CCC"], fewshot_rows
    )
    endpoint_definition_table = md_table(["中文名称", "数据字段", "单位", "当前用途"], endpoint_definition_rows)

    axis_names = {
        "deformation_compliance": "变形-柔顺性",
        "flesh_resistance_energy": "果肉阻力-能量",
        "skin_rupture_resistance": "表皮破裂阻力",
    }
    axis_rows = []
    for target in ["deformation_compliance", "flesh_resistance_energy", "skin_rupture_resistance"]:
        p0 = dataframe_lookup(axis, target, model="PLSR")
        c0 = dataframe_lookup(axis, target, model="1D-CNN")
        p5 = dataframe_lookup(axis_pls_fewshot, target, shots=5)
        p20 = dataframe_lookup(axis_pls_fewshot, target, shots=20)
        c5 = dataframe_lookup(axis_cnn_fewshot, target, shots=5)
        c20 = dataframe_lookup(axis_cnn_fewshot, target, shots=20)
        axis_rows.append(
            [
                axis_names[target],
                fmt(p0.r2),
                fmt(c0.r2),
                fmt(p5.r2_mean),
                fmt(c5.r2_mean),
                fmt(p20.r2_mean),
                fmt(c20.r2_mean),
            ]
        )
    axis_table = md_table(
        ["机械轴", "0-shot PLSR", "0-shot CNN", "5-shot PLSR", "5-shot CNN", "20-shot PLSR", "20-shot CNN"],
        axis_rows,
    )

    sensitivity_pivot = sensitivity.pivot(index="target_label", columns="cohort", values="r2")
    sensitivity_rows = []
    for target in TARGET_ORDER:
        label_en = validity.loc[validity["target"].eq(target), "target_label"].iloc[0]
        row = sensitivity_pivot.loc[label_en]
        sensitivity_rows.append(
            [
                LABELS_ZH[target],
                fmt(row["Hard-valid sensitivity"]),
                fmt(row["High-confidence analysis"]),
                fmt(row["Strict 10%"]),
            ]
        )
    sensitivity_table = md_table(["端点", "硬有效R2", "高置信主分析R2", "严格10% R2"], sensitivity_rows)

    claim_table = md_table(
        ["当前/潜在表述", "证据状态", "主要风险", "建议改写"],
        [
            ["质构数据质量差", "不支持", "忽略ICC约0.91和11,004条可解析曲线", "测量总体可靠，但存在端点、批次和域差异"],
            ["模型实现新品种个体预测", "证据不足", "品种内中心化R2最高仅0.119", "模型可预测部分总体差异；个体排序能力仍弱"],
            ["5-shot恢复个体预测", "过强", "截距校准不改变品种内排序或Pearson r", "5-shot主要校正新品种均值偏差"],
            ["发现品种效应", "部分支持", "14/16品种仅一个批次，品种与批次混杂", "发现品种/批次复合域偏移"],
            ["删除10%坏果提高模型", "不支持", "严格队列多数LOCO结果变差", "严格队列仅作发布与敏感性分析"],
            ["CNN优于化学计量学", "不支持", "主分析队列CNN整体不及PLSR", "CNN是阴性容量对照"],
            ["预测口感和贮藏期", "无直接标签", "没有感官或货架期数据", "预测与商品品质相关的机械代理表型"],
            ["pH代表含酸量", "不成立", "pH不是可滴定酸", "严格称pH，除非找到独立酸度数据"],
            ["严格QC规则已预注册", "档案未证明", "QC层级在敏感性结果后调整", "称为透明的迭代式、模型无关QC并完整报告版本史"],
        ],
    )

    strategy_table = md_table(
        ["论文方向", "当前证据强度", "优势", "致命短板", "建议"],
        [
            ["A. 大规模机械表型资源", "较强", "5,502果、11,004曲线、双重复、16标签", "尚缺公开许可、外部数据基准和完整方法验证", "优先考虑"],
            ["B. 跨域失败与验证方法学", "中等偏强", "随机拆分与LOCO/LOBO差异鲜明", "品种和批次混杂，需更严谨分解", "可与A合并"],
            ["C. 少样本部署", "当前偏弱", "均值校准显著改善汇总指标", "截距不改善个体排序", "需升级为斜率/表示层校准并嵌套验证"],
            ["D. 深度学习优越性", "弱", "已有CNN/Transformer尝试", "主分析无优势、CNN仅单种子", "不宜作为当前主线"],
            ["E. 口感/贮藏商品品质", "概念潜力高", "机械表型具有生物学关联", "缺少感官与贮藏标签", "作为后续实验，不作当前直接结论"],
        ],
    )

    file_table = md_table(
        ["内容", "相对项目根路径", "状态"],
        [
            ["主QC账本", "data/processed/texture_qc/texture_qc_ledger.parquet", "5,502果、三套队列标志"],
            ["严格英文发布集", "../NIRs_plums_research_ready_en_v1.2.0", f"独立审计{audit['status']}"],
            ["质构端点", "data/processed/texture/texture_sample_features.parquet", "双重复派生特征"],
            ["PLSR留一品种预测", "results/models/texture_pls_loco_analysis", "16外层折"],
            ["CNN留一品种预测", "results/models/texture_cnn_loco_analysis", "单随机种子"],
            ["少样本结果", "results/models/texture_fewshot_analysis", "100次重复；截距校准"],
            ["批次结果", "results/models/texture_pls_multibatch", "仅Konglongdan与Weiwang"],
            ["当前英文稿", "manuscript/manuscript_texture_first_revision.md", "需按本报告红队问题重写"],
            ["重现脚本", "src/", "关键脚本已通过py_compile"],
        ],
    )

    figure_curve = (project / "results" / "texture_atlas" / "figures" / "fig_texture_curve_mechanics.png").as_posix()
    figure_atlas = (project / "results" / "texture_atlas" / "figures" / "fig_texture_phenotype_atlas.png").as_posix()
    figure_qc = (project / "results" / "texture_qc" / "figures" / "fig_texture_qc_audit.png").as_posix()
    figure_prediction = (project / "results" / "texture_prediction_analysis" / "fig_texture_prediction_transfer.png").as_posix()
    figure_critical = (project / "results" / "project_review" / "fig_critical_validity_audit.png").as_posix()
    figure_concept = (project / "manuscript" / "figures" / "concept" / "nir_texture_graphical_abstract_base.png").as_posix()

    report = f"""# NIRs_plum 无损检测项目独立审阅报告

## 当前版本、证据边界与论文提升议程

**文档性质：** 供其他线程、统计学/光谱学/采后生物学研究者进行独立红队审阅  
**项目根目录：** `{project}`  
**数据发布版本：** `NIRs_plums_research_ready_en_v1.2.0`  
**QC版本：** `0.2.0`  
**报告日期：** 2026-08-06  
**状态：** 当前稿不建议直接投稿；需先完成关键有效性重审与叙事重构

> [总判断] 该项目最强资产是规模巨大、双重复且可追溯的机械质构曲线资源；最弱环节是把汇总R2和截距校准改善解释为新品种内个体果实预测。当前证据支持“测量可靠、存在光谱关联、存在显著域偏移”，但尚不足以支持“已实现稳健的新品种个体质构预测”。

<!-- PAGE_BREAK -->

# 1. 报告目的与审阅者使用方式

本报告不是成果宣传稿，而是完整暴露当前项目的设计、数据、处理规则、模型结果、失败结果和可能的论证漏洞。用户将把它交给其他线程或其他研究人员审阅；审阅意见返回后，再由当前执行者据此进行下一轮优化。因此，本报告有意区分“已经完成”“目前观察到”“可以防御的结论”“仍需验证”和“当前不应声称”五类内容。

审阅者不需要接受当前论文方向。请优先判断三个根本问题：第一，这批数据最适合写成数据/表型资源论文、跨域方法论文，还是应用预测论文；第二，现有验证指标是否真的衡量了单果层面的预测；第三，品种、批次、操作员和成熟度混杂是否使当前“品种迁移”叙事无法成立。

建议审阅顺序：先读第2节执行摘要和第10节红队审计，再看第9节完整数字，最后根据第13节问题清单提交意见。报告末尾提供了统一反馈模板。

# 2. 执行摘要

## 2.1 已完成的核心工作

- 将混乱命名的光谱、果实性状和ARC文件匹配到5,502个果实级身份，统一英文品种名与标准样本ID。
- 解析11,004条质构仪ARC曲线，提取双重复机械特征；仅1条曲线发生硬性失败。
- 建立模型无关的QC账本，保留高置信主分析、严格10%发布和硬有效敏感性三套队列。
- 建立4,941果的严格完整英文发布集，包含16个品种/选系、19批次、21,130个测量文件和228波长点NIR矩阵。
- 完成随机果实拆分、嵌套留一品种PLSR、1D-CNN、机械轴、少样本截距校准和有限的留一批次验证。
- 生成质构曲线图、表型图谱、QC审计图、迁移验证图和概念图，并形成一版质构优先英文论文草稿。

## 2.2 当前最可信的结果

- 九个主要质构端点完整队列中位ICC(A,1)={qc_summary['median_icc_full']:.3f}，高置信主分析队列为{qc_summary['median_icc_high_confidence_analysis']:.3f}。质构测量总体不是“非常不理想”。
- 三个描述性机械轴解释95.1%的标准化端点变异，品种/域间机械差异明显。
- 品种分层随机拆分PLSR R2为0.235-0.479，说明光谱中存在与质构相关的信息。
- 嵌套留一品种汇总PLSR R2为-0.103至0.338，显示未见品种上的迁移显著下降。
- 严格删除约10%后多数LOCO结果更差，因此不能把全部被删果实称为测量错误。

## 2.3 当前最严重的问题

- 品种内中心化R2仅为{critical['within_cultivar_centered_r2_range'][0]:.3f}至{critical['within_cultivar_centered_r2_range'][1]:.3f}；16个留出品种内Pearson r中位数仅0.090至0.343。这表明汇总R2很大一部分来自域间均值和尺度差，而非新品种内部的可靠单果排序。
- 5-shot方法只估计一个截距。它能校正新品种均值，却在数学上不能改变该品种内部预测排序、Pearson r或中心化误差结构。因此“少样本恢复个体预测”是过强表述。
- 16个品种/选系中14个只有1个记录批次；只有Konglongdan有3批、Weiwang有2批。所谓“品种偏移”与日期、操作员、成熟度和批次协议高度混杂。
- 当前CNN只运行一个种子；Transformer只在严格队列运行；研究计划中的SVR、树模型、分层模型、域适配和公平重复比较尚未完成。
- 没有感官评价、贮藏时间、腐烂、失水或货架期标签，不能直接宣称预测口感或贮藏期。

## 2.4 报告给出的初步路线

当前最值得保留的是“大规模机械表型资源 + 对常规随机验证的系统性警示”。如果继续主打预测，应首先重建评价框架，把三个不同任务分开：已知品种内个体排序、新品种均值校准、新品种内个体排序。只有第三项达到可接受水平，才能把少样本方法称为个体预测恢复。

# 3. 科学问题与原始实验逻辑

## 3.1 实际测量顺序

原始操作顺序为：完整果实 -> NIR光谱 -> 单果重 -> 质构仪 -> 可溶性固形物 -> pH。NIR是在果实完整状态下获取的无损输入；后续质量、质构、SSC和pH都是参考表型。质构仪虽对果实产生穿刺，但其科学角色与测糖和pH相同：为无损光谱模型提供表型标签，而不是作为模型部署时的输入。

该逻辑使“直接由完整果实NIR预测机械质构”成为合理问题。表皮破裂、果肉阻力、变形、机械功和黏附与口感、成熟、抗损伤和采后变化具有生物学关系，但本项目只能称这些指标为相关机械代理表型。

## 3.2 当前论文拟解决的问题

1. ARC双重复曲线能否形成可靠、可解释的机械质构端点？
2. NIR是否包含这些端点的可预测信息？
3. 随机果实拆分与完全未见品种验证之间有多大差距？
4. 少量新品种参考果实能否改善迁移？
5. 操作/批次变化在多大程度上限制模型？

第1和第2问已有肯定但有限的证据；第3问证据较强；第4问目前只证明均值偏差可校正，尚未证明个体排序恢复；第5问受批次重复不足限制。

## 3.3 Akagi桃研究的借鉴范围

重点参考Masuda等（2023，Postharvest Biology and Technology 201:112348）的多性状桃果实研究：从样本规模、性状分布、相关性、建模流程、连续回归到解释性可视化的叙事顺序。该研究使用1,521个Hakuho桃和7个性状，并比较多个CNN。当前李子项目数据规模更大、具有质构双重复和多品种，但域混杂也更严重。借鉴应限于研究结构和图形叙事，不应暗示两套验证设计或任务难度等价。

# 4. 数据工程、身份匹配与发布集

## 4.1 命名标准化

原始数据中存在中文名、拼音、英文缩写、批次代号和同品种多别名，例如味王/WW、味帝/WD。整理后，每个果实具有稳定的英文标准ID，格式为`plum-<cultivar-slug>-b<batch>-f<fruit>`；交叉表保留原始样本ID、原始文件路径、标准路径和SHA-256。

名称标准化解决的是身份可追溯性，不等同于证明所有匹配都无误。高风险身份应由审阅者抽查：多批次品种、编号重置处、源文件日期切换处和原始命名相似的味系品种。

## 4.2 三层队列与各品种数量

{cultivar_table}

主分析保留5,430果，只排除72个高置信技术异常。严格质构队列保留4,952果；进一步要求质量、SSC、pH、公共`c`光谱和两条有效ARC曲线后，英文发布集为4,941果。硬有效敏感性队列保留5,500果。

## 4.3 发布集完整性

独立审计状态为**{audit['status']}**：验证{audit['checksum_entries_verified']:,}项校验和、{audit['measurement_files']:,}个测量文件、{audit['samples']:,}个唯一果实；没有非法标准ID或不完整样本。NIR矩阵为{audit['nir_matrix_shape'][0]:,} x {audit['nir_matrix_shape'][1]}，有限值比例为{audit['finite_nir_fraction']:.1f}。

公共建模光谱为`c`扫描。严格发布集全部4,941果具有`c`扫描；其中味帝683果另有`t`扫描。`t`扫描不是跨品种公共模态，当前主模型未使用。审阅者应确认`c`和`t`的仪器含义、采集位置与协议，否则不能把未使用`t`简单解释为重复测量。

## 4.4 常规性状语义

发布集含单果重（g）、可溶性固形物（%）和pH。pH不是可滴定酸含量；除非找到独立酸度测定文件，正文不得写“含酸量”。

# 5. ARC解析与机械端点

## 5.1 曲线处理

11,004条ARC曲线被解码为时间、力、位置和状态通道，采样率约400 Hz。基线取前0.20 s中位数，噪声用稳健MAD估计，接触点由持续超过阈值的力识别，加载终点由位置反转确定。加载峰用于表皮破裂，破裂后阶段用于果肉阻力和功，回程负力用于黏附。

当前力值乘以1,000以匹配存档预览中的克力显示，但尚无独立仪器证书或导出说明验证该缩放。位置通道无法确认物理单位，所有位移和功保持`raw_position_unit`，不得改写为mm或mJ。

## 5.2 九个主要端点

{endpoint_definition_table}

峰数重复性较差，当前不作为主要预测端点；最大加载力与表皮破裂力高度冗余，也不单独作为主端点。端点选择同时参考机械意义、重复性和冗余，但这一选择是在当前数据上完成，属于探索性特征注册，不是外部预注册。

![图1 机械可解释质构曲线、典型双重复、异常双重复和5,430果一致性]({figure_curve})

# 6. QC体系与删除策略

## 6.1 QC证据

QC版本0.2.0不使用任何预测残差。证据分为五类：曲线采集稳定性、双重复不一致、批次内多变量质构极端、NIR采集/PCA异常、常规性状极端。技术证据是排除的必要条件；单纯生物学极端不能独立删除果实。

“中等异常”使用全数据集内第80百分位阈值；严重阈值包括稳健z、异常特征数量、基线噪声比例和硬失败。高置信主分析仅删除硬失败，或严重技术证据且至少有一类独立支持。严格队列进一步删除多类中等证据一致的果实。

## 6.2 QC结果

- 完整匹配：5,502果。
- 硬有效：5,500果。
- 高置信主分析：5,430果，排除72果（1.3%）。
- 严格质构：4,952果，排除550果（10.0%）。
- 严格完整英文发布：4,941果。
- 严格队列批次排除率范围：{100*qc_summary['batch_exclusion_rate_range'][0]:.1f}%-{100*qc_summary['batch_exclusion_rate_range'][1]:.1f}%。
- 主分析保留原始IQR中位比例{100*qc_summary['median_analysis_iqr_retained_fraction']:.2f}%，保留2%-98%范围{100*qc_summary['median_analysis_2_98_range_retained_fraction']:.2f}%。

## 6.3 必须透明说明的时间顺序

严格10%队列完成后发现多数LOCO性能下降，随后项目把72个高置信排除的5,430果队列作为主分析。这一调整具有科学理由，但当前档案不能证明它在查看性能前已前瞻性预注册。因此稿件不应使用“predeclared tiers”一类措辞。应明确写为：模型无关QC规则经过版本化迭代，全部队列和敏感性结果均保留；主分析选择基于避免过度删除，而非宣称该选择此前已注册。

![图2 QC队列流、证据组合、重复性、批次保留和性状范围]({figure_qc})

# 7. 描述性机械表型结果

九端点完整队列中位ICC(A,1)=0.907，高置信主分析=0.910。加载刚度ICC约0.982，破裂后功约0.926，表皮破裂力约0.910；原始位移6处力较低但仍约0.779。严格删除10%只带来很小的ICC变化，说明原始测量活动总体成功。

品种间端点分布和机械指纹差异明显。三个描述性轴解释95.1%变异：果肉阻力-能量、变形-柔顺性、表皮破裂阻力。预测轴在每个留一品种折中用训练品种的中位数和IQR重新缩放，避免直接使用外层测试品种的目标分布。

需要注意：描述性PCA/varimax轴的解释和预测轴的固定端点组合不是同一个统计对象。当前轴名称具有合理机械含义，但仍需要外部专家审查组合权重和方向，避免为了提升预测而事后定义复合终点。

![图3 质构样本规模、分布、相关图谱、品种指纹和三轴表型空间]({figure_atlas})

# 8. 当前建模与验证设计

## 8.1 随机果实拆分

高置信队列进行5次品种分层80:20随机拆分。外层训练内通过4折品种分层果实CV选择预处理和PLS成分。因为训练和测试均包含相同品种，该设计衡量熟悉域内插值，不代表新品种部署。

## 8.2 嵌套留一品种PLSR

外层依次完全留出16个品种/选系；内层GroupKFold按剩余品种选择raw、SNV、一阶SG导数、SNV+导数，以及4/8/12/16/24个PLS成分。该设计避免外层品种参与超参数选择，是当前最公平的主基线。

## 8.3 1D-CNN与Transformer

1D-CNN用raw、SNV和导数三通道联合预测九端点，使用品种平衡训练、Huber损失、AdamW、早停和独立验证品种。主分析队列目前仅完成一个随机种子。Transformer只在严格队列运行且多数端点为负，不能与主分析PLSR/CNN做完全公平比较。

## 8.4 少样本校准

每个外层留出品种随机抽取k个果实，计算`mean(y_true - y_pred)`作为单一截距，再从评价集中移除这些果实；每个k重复100次。该程序无果实重叠泄漏，但它仅校正均值，不学习斜率、排序、非线性或表示。因此适合称“target-domain intercept calibration”，不宜笼统称few-shot learning。

## 8.5 批次验证

只有Konglongdan和Weiwang存在多批次，形成5个留一批次测试。其余14个标签无法把品种与批次分离。批次留出PLSR九端点R2均为负，范围-0.564至-0.090；20-shot截距校准后的九端点中位R2约0.031。

# 9. 完整数值结果

## 9.1 测量可靠性、随机插值与跨品种迁移

{endpoint_table}

最重要的对比不是随机R2和LOCO汇总R2本身，而是LOCO汇总R2与品种内中心化R2的差异。破裂位移汇总R2=0.338，但去除各留出品种均值后只有0.119；加载刚度汇总R2=0.218，中心化后为-0.112。多数端点的品种内中心化R2为负。

## 9.2 截距校准结果

{fewshot_table}

这些数字说明少量参考果实可以估计新品种均值偏差。它们不说明单果排序得到改善，因为添加常数不会改变同一品种内部的Pearson r。审阅者应决定：如果实际应用只需要估计一个新品种总体质构水平，该结果可能有价值；如果目标是分选同一品种内不同果实，当前结果明显不足。

## 9.3 机械轴

{axis_table}

机械轴的汇总R2高于多数单端点，但仍应补充品种内中心化和排名指标。5-shot轴改善同样主要来自截距校正。不能仅凭轴R2=0.678就宣称5个果实使模型恢复了新品种内部的机械分选能力。

## 9.4 QC敏感性

{sensitivity_table}

严格10%队列中破裂位移R2从主分析0.338降至0.138；表皮破裂力从0.076降至-0.098。硬有效5,500果队列多项结果与主分析接近或更好。这是反对“删得越多越干净、模型越好”的直接证据。

## 9.5 PLSR与CNN

主分析队列中PLSR在九端点汇总R2上均不低于CNN。PLSR最佳为破裂位移0.338，CNN为0.248；PLSR机械轴为0.463/0.261/0.226，CNN为0.353/0.120/0.119。当前结果不支持深度模型优越性。

![图4 当前验证层级、随机-迁移差距、模型比较和少样本汇总结果]({figure_prediction})

# 10. 红队有效性审计：当前稿件为什么还不够好

## 10.1 汇总R2掩盖品种内个体拟合不足

汇总所有留出品种后计算R2，会同时利用不同品种之间的均值和尺度差异。对实际分选而言，更关键的是模型能否在一个未知品种内部正确排序不同果实。当前品种内中心化R2范围为-0.231至0.119，只有破裂位移为明显正值但仍偏低。

## 10.2 5-shot改善主要是均值校正

截距校准后force drop等端点汇总R2大幅增加，是因为消除了品种均值偏差。任何加常数操作都不能改变该品种果实的排序。当前英文稿中“restore useful prediction”“convert fragile universal models into useful predictors”等表述需要降级，除非后续斜率校准、局部重训练或域适配真正提高品种内评价。

## 10.3 品种与批次严重混杂

14/16标签只有一个批次。若某个品种由不同人员、日期、成熟度或协议测量，那么LOCO误差不能被纯粹归因于遗传品种。当前标题中的“cultivar shift”应暂改为“cultivar/batch domain shift”，或通过多批次品种分析证明遗传和操作域可以分离。

## 10.4 QC主队列选择具有事后成分

高置信队列的规则不使用模型残差，这是优点；但把它升级为主分析发生在看到严格队列性能下降之后。完整报告这一历史是可以接受的，声称前瞻性预注册则不可接受。下一轮分析应冻结QC 0.2.0，不再依据新模型结果改变样本。

## 10.5 模型比较尚不充分

- CNN只有一个随机种子，不能评价训练方差。
- Transformer未在相同主队列和相同外层折完整重跑。
- 尚缺SVR、核方法、梯度提升、岭回归、随机森林和简单光谱距离基线。
- 尚缺层级/混合效应模型，将品种均值和个体偏差分开建模。
- 未报告最简单的品种均值、训练集全局均值和仅质量/SSC/pH基线。
- XAI和波段稳定性尚未完成；在性能评价未稳定前做漂亮的解释图风险较高。

## 10.6 端点和物理单位仍需仪器专家确认

力缩放1,000、接触点阈值、平滑、固定原始位移6、加载段/卸载段界定均需由熟悉仪器的人审阅。位置单位不明限制机械功的物理解释。建议用原仪器软件对20-30条曲线逐项核对派生值。

## 10.7 生物学锚点不足

没有硬度计对照、感官评价、成熟度等级、储藏天数、失水、腐烂或货架期。机械指标本身有价值，但当前稿无法证明它们对消费者口感或实际贮藏终点的解释度。

![图5 红队审计：汇总R2、品种内拟合、截距校准和品种-批次混杂]({figure_critical})

# 11. 可防御结论与高风险表述

{claim_table}

当前最安全的中心结论是：大规模双重复ARC档案能产生可靠且多维的李子机械表型；NIR在已知品种混合样本中包含质构相关信号，但模型向新域迁移时显著下降；现有截距校准能校正新品种均值，而新品种内部的单果排序仍是未解决问题。

# 12. 论文方向选择与提升路线

## 12.1 可选论文方向

{strategy_table}

## 12.2 优先级P0：在任何新模型前完成

1. 冻结QC 0.2.0和三套队列，不再根据预测性能更改删除规则。
2. 重写标题、摘要和结论，加入品种/批次混杂和品种内中心化结果。
3. 将评价明确拆成：熟悉品种内插值、未知域均值、未知域内排序。
4. 对所有端点报告汇总R2、品种内中心化R2、品种宏平均R2、每品种Pearson/Spearman、CCC和RMSE。
5. 加入全局均值、品种均值oracle、质量/SSC/pH和简单线性基线，判断NIR的独立增益。
6. 删除“predeclared”表述，增加QC版本史和分析决策时间线。

## 12.3 优先级P1：公平重建模型比较

1. 在完全相同的16个外层折和训练内超参数选择下比较PLSR、ridge/elastic net、SVR、树模型和1D-CNN。
2. 深度模型至少5个种子，报告均值、标准差和每折分布；任何早停验证品种必须与外层测试严格隔离。
3. 建立层级模型：分别预测域均值与品种内偏差，避免一个汇总R2同时混合两种任务。
4. 比较截距、截距+斜率、少量PLS更新、局部加权校准和表示层微调；所有校准果实从评价集中移除，并在每个品种内重复抽样。
5. 尝试严格训练折内的光谱标准化迁移、CORAL/域对齐或批次稳健损失，但必须与简单基线比较，避免域适配包装弱结果。
6. 评估多任务学习：SSC、pH、质量作为辅助目标而非部署输入，检验是否形成更稳定的共享表示。

## 12.4 优先级P2：机械表型和生物学加强

1. 请仪器专家盲审ARC解析规则，并用原软件手工复核代表曲线。
2. 分析两次穿刺的系统位置差异，而不仅是取均值；位置差异本身可能是果实异质性表型。
3. 对原曲线做功能主成分、形状距离或自监督曲线表示，判断九个手工端点是否丢失信息。
4. 若存在原始成熟等级、采收日期、果皮颜色、硬度计或储藏记录，重新纳入身份匹配。
5. 新一轮采集应安排跨操作员控制果、同品种跨批次、仪器日内/日间标准和完整元数据。

## 12.5 优先级P3：论文与图形

1. 完成系统性文献检索，比较水果NIR硬度/质构、跨品种迁移、校准转移和机械曲线建模。
2. 主图首先展示任务分解与验证诚实性，而不是只展示样本量和高R2。
3. 将品种内中心化性能和每品种分布放入主图；少样本图明确标注“intercept-only”。
4. XAI只在稳定模型和稳定端点上进行，并报告波段在外层折/种子间的一致性。
5. Image2概念图只作视觉底图，所有文字、数字和箭头用可编辑矢量层叠加。

# 13. 提交给审阅者的具体问题

## 13.1 科学定位

1. 这批数据最强的论文类型是什么：资源、方法学警示、预测应用，还是需要拆成两篇？
2. 九个端点中哪些具有最清晰的机械/采后意义，哪些应降为补充？
3. 三个机械轴是否有足够理论基础，还是应改用数据驱动且外层折内拟合的潜变量？
4. 在没有感官和储藏标签时，如何准确描述商品品质价值？

## 13.2 统计与验证

5. 汇总LOCO R2是否适合作为主指标？推荐的主指标组合是什么？
6. 品种内中心化R2的定义是否合理，是否应改为分层R2、条件R2或rank-based评价？
7. 5-shot截距校准应如何命名和定位？
8. 如何在仅两个多批次品种的条件下处理品种/批次混杂？
9. 高置信QC队列作为主分析是否可接受？需要怎样的敏感性或决策记录？
10. 当前ICC使用绝对一致性ICC(A,1)是否合适？是否需要Bland-Altman、异方差或位置效应分析？

## 13.3 光谱与模型

11. 901-1701 nm、228点的预处理候选是否足够？
12. 是否需要散射校正、OSC、波段选择、区间PLS或非线性核方法？
13. 如何设计公平的深度学习和PLSR比较，包括种子、超参数预算和内层验证？
14. 多任务SSC/pH/质量是否有可能提升质构表示，还是会引入捷径？
15. 哪一种域适配/校准转移方法最适合当前批次结构？
16. 是否应直接从曲线表型潜变量或原曲线嵌入建模，而不是九个端点？

## 13.4 仪器与采后生物学

17. 力乘1,000和接触/破裂定义是否符合原仪器语义？
18. `raw_position_unit`能否从仪器配置、探头速度或原软件恢复为物理位移？
19. 两个穿刺位置应取均值、最大值、最小值，还是建模位置异质性？
20. 哪些端点最可能与口感、耐储和损伤相关，但又不造成过度因果推断？

## 13.5 论文与展示

21. 当前题目是否过度强调few-shot“enables prediction”？
22. 最有说服力的3-5张主图应如何排列？
23. 哪些数字必须进入摘要，哪些应只放补充材料？
24. Akagi桃研究的哪些方法可以合理复现，哪些不适用于本项目？
25. 以当前证据水平，合理的目标期刊层级和补强要求是什么？

# 14. 审阅意见返回模板

请每位审阅者按下列格式返回，便于下一轮统一实施：

1. **总体判断：** 当前论文最合适的主线和一句话中心结论。
2. **不可接受的问题：** 最多列5项，说明为什么会导致拒稿。
3. **必须重做的分析：** 给出输入、方法、验证方式、主指标和成功判据。
4. **建议删除的主张/图表：** 指出具体段落或图。
5. **建议新增的主张/图表：** 说明所需证据。
6. **端点优先级：** 九端点和三机械轴中建议保留的主终点。
7. **论文方向：** A资源 / B域偏移方法 / C少样本部署 / D其他。
8. **优先级：** P0阻断投稿、P1显著提高、P2锦上添花。
9. **可执行细节：** 尽量提供算法、参数范围、对照基线或参考文献。
10. **目标期刊建议：** 说明达到该期刊前仍缺什么。

# 15. 文件导航与可复现性

{file_table}

所有原始来源文件未被删除。严格发布集包含标准英文文件名、样本交叉表、原始相对路径、SHA-256和独立审计报告。当前报告新增的红队评价表位于`results/project_review/`，生成脚本为`src/build_project_review_metrics.py`和`src/build_independent_review_report.py`。

# 16. 当前图形资产评价

质构曲线图已经能清楚解释加载、破裂、果肉阻力、卸载黏附和双重复；表型图谱能显示数据规模与品种差异；QC图能说明删除证据和范围保留；预测图能显示随机-LOCO差距。但当前主图仍不够突出“汇总分数与品种内排序的分离”，因此本报告新增图5应成为下一版论文结构的重要候选。

Image2生成的概念图适合作为图形摘要底图，但不能代替数据驱动图，也不应把“口感/贮藏”画成已被直接预测的标签。

![图6 Image2生成的NIR到多维机械质构概念底图]({figure_concept})

# 附录A. 关键指标解释

- **ICC(A,1)：** 两次穿刺的绝对一致性单次测量可靠性；高ICC说明重复性好，不保证NIR可预测。
- **随机拆分R2：** 同品种同时存在于训练和测试，衡量熟悉域内插值。
- **LOCO汇总R2：** 每次留一品种，合并所有外层预测后计算；同时受品种间和品种内差异影响。
- **品种内中心化R2：** 对每个留出品种分别去除真实和预测均值后计算，聚焦品种内个体偏差拟合。
- **品种宏平均R2：** 每个品种单独计算R2后等权平均；对小品种和均值偏差非常敏感。
- **CCC：** 同时衡量相关性和均值/尺度一致性；截距校准可显著提高CCC。
- **5-shot截距校准：** 用5个有标签果实估计均值残差；不改变同一品种内部排序。

# 附录B. 已知边界与缺失元数据

- 无操作员身份，因此不能把批次差异归因于具体人员。
- 大多数品种只有一个批次，不能独立估计遗传品种与批次效应。
- 无果园、树号、季节、采收日、成熟等级等完整结构化元数据。
- 无感官、贮藏时间、腐烂、失水和货架期终点。
- 位置物理单位未确认；力缩放仍需仪器资料核验。
- 仅味帝具有额外`t`扫描，公共主模型仅使用`c`扫描。
- 数据许可未指定，公开发布前必须解决授权和隐私/来源问题。

# 附录C. 报告结论

本项目值得继续，但必须改变“用更漂亮的图和更复杂的模型包装当前R2”的思路。下一轮优化应首先修复任务定义和评价：把域均值校准与品种内个体预测分开，诚实处理品种/批次混杂，再决定模型复杂度。只有当新品种内排序或中心化R2得到真实提高，few-shot才能成为论文亮点；否则，最有价值的论文是大规模机械表型资源及其对不严谨随机验证的警示。
"""
    return report


NAVY = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(35, 47, 62)
MUTED = RGBColor(95, 105, 115)
LIGHT_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None, italic: bool | None = None, mono: bool = False) -> None:
    western = "Consolas" if mono else "Calibri"
    east_asia = "Microsoft YaHei"
    run.font.name = western
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), western)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), western)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def column_widths(n: int) -> list[int]:
    if n == 2:
        return [2600, 6760]
    if n == 3:
        return [2300, 3530, 3530]
    if n == 4:
        return [2200, 2386, 2387, 2387]
    first = 2100
    remaining = 9360 - first
    base = remaining // (n - 1)
    widths = [first] + [base] * (n - 1)
    widths[-1] += 9360 - sum(widths)
    return widths


def add_inline(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    for token in filter(None, pattern.split(text)):
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=9, mono=True, color=DARK)
        else:
            run = paragraph.add_run(token)
            set_run_font(run)


def add_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        cells = [cell.strip().replace("\\|", "|") for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    rows = [rows[0]] + rows[2:]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        tr_pr = table.rows[r_idx]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline(paragraph, value)
            for run in paragraph.runs:
                set_run_font(run, size=7.4 if len(rows[0]) >= 6 else 8.2, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                repeat = OxmlElement("w:tblHeader")
                repeat.set(qn("w:val"), "true")
                tr_pr.append(repeat)
            elif c_idx > 1 or (len(rows[0]) <= 4 and c_idx > 0):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, column_widths(len(rows[0])))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    add_inline(p, text)
    for run in p.runs:
        set_run_font(run, size=10.2, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def set_headers_footers(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("NIRs_plum | Independent Scientific Review Packet")
        set_run_font(run, size=8.5, color=MUTED, bold=True)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_after = Pt(0)
        add_page_field(fp)


def add_cover(doc: Document) -> None:
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(18)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("INDEPENDENT SCIENTIFIC REVIEW PACKET")
    set_run_font(run, size=10, color=BLUE, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("NIRs_plum 无损检测项目")
    set_run_font(run, size=28, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    run = subtitle.add_run("当前证据、关键缺陷与论文提升议程")
    set_run_font(run, size=16, color=DARK)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(28)
    run = meta.add_run("Version 1.0 | 2026-08-06 | For independent review")
    set_run_font(run, size=10, color=MUTED, italic=True)
    add_callout(
        doc,
        "当前状态：不建议直接投稿。最强资产是5,502果、11,004条双重复ARC曲线；最大风险是把域均值校准误解为新品种内部的单果预测能力。",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run("Prepared for statistical, spectroscopy, postharvest and manuscript red-team review")
    set_run_font(r, size=9.5, color=MUTED)
    doc.add_page_break()


def prepare_docx_image(source: Path, asset_dir: Path, max_width: int = 1200) -> Path:
    asset_dir.mkdir(parents=True, exist_ok=True)
    target = asset_dir / f"{source.stem}__docx.jpg"
    with Image.open(source) as image:
        image = image.convert("RGB")
        if image.width > max_width:
            height = round(image.height * max_width / image.width)
            image = image.resize((max_width, height), Image.Resampling.LANCZOS)
        image.save(target, format="JPEG", quality=82, optimize=True, progressive=False, dpi=(150, 150))
    return target


def markdown_to_docx(markdown: str, output: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped == "<!-- PAGE_BREAK -->":
            doc.add_page_break()
            index += 1
            continue
        if stripped.startswith("| "):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_table(doc, table_lines)
            continue
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)$", stripped)
        if image_match:
            caption, path = image_match.groups()
            if "nir_texture_graphical_abstract_base" in path:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(8)
                run = paragraph.add_run(f"{caption}。完整图像见项目 manuscript/figures/concept/ 目录。")
                set_run_font(run, size=9, color=MUTED, italic=True)
                index += 1
                continue
            docx_image = prepare_docx_image(Path(path), output.parent / "_docx_assets")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            run = p.add_run()
            inline_shape = run.add_picture(str(docx_image), width=Inches(6.35))
            inline_shape._inline.docPr.set("descr", caption)
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_before = Pt(2)
            cap.paragraph_format.space_after = Pt(10)
            cap.paragraph_format.keep_with_next = False
            r = cap.add_run(caption)
            set_run_font(r, size=8.5, color=MUTED, italic=True)
            index += 1
            continue
        if stripped.startswith("> "):
            add_callout(doc, stripped[2:].strip())
            index += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            add_inline(paragraph, text)
            index += 1
            continue
        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:])
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            add_inline(paragraph, numbered.group(1))
            index += 1
            continue
        paragraph = doc.add_paragraph()
        add_inline(paragraph, stripped)
        index += 1

    set_headers_footers(doc)
    doc.core_properties.title = "NIRs_plum 无损检测项目独立审阅报告"
    doc.core_properties.subject = "Current evidence, validity risks, and paper improvement agenda"
    doc.core_properties.author = "Project analysis team"
    doc.core_properties.keywords = "NIR, plum, texture, spectroscopy, independent review"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--project-root", type=Path, required=True)
    parser.add_argument("--markdown-output", type=Path, required=True)
    parser.add_argument("--docx-output", type=Path, required=True)
    args = parser.parse_args()
    project = args.project_root.resolve()
    markdown = build_markdown(project)
    args.markdown_output.resolve().write_text(markdown, encoding="utf-8")
    markdown_to_docx(markdown, args.docx_output.resolve())
    print(json.dumps({"markdown": str(args.markdown_output.resolve()), "docx": str(args.docx_output.resolve())}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
