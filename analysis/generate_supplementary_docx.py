from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from generate_manuscript_docx import (
    DARK_BLUE,
    LIGHT_FILL,
    configure_section,
    configure_styles,
    set_cell_shading,
    set_repeat_table_header,
    set_run,
    set_table_geometry,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "manuscript/Supplementary_material_plum_NIR.docx"


def add_table(doc: Document, frame: pd.DataFrame, caption: str, widths: list[int], note: str) -> None:
    paragraph = doc.add_paragraph(style="Table Caption")
    set_run(paragraph.add_run(caption), size=9.5, bold=True, color=DARK_BLUE)
    table = doc.add_table(rows=1, cols=len(frame.columns))
    set_repeat_table_header(table.rows[0])
    for index, name in enumerate(frame.columns):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(str(name)), size=7.2, bold=True, color=DARK_BLUE)
    for row in frame.fillna("—").itertuples(index=False, name=None):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            p = cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index < 2 else WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(str(value)), size=7.0)
    set_table_geometry(table, widths)
    p = doc.add_paragraph(style="Table Note")
    set_run(p.add_run(f"Note: {note}"), size=8.5, italic=True)


def main() -> None:
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    set_run(p.add_run("Supplementary material"), size=22, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(
        p.add_run("Cultivar shift constrains universal near-infrared phenotyping of plum quality but few-shot hierarchical calibration restores transferability"),
        size=13,
        bold=True,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("[AUTHOR NAMES]"), size=11)
    doc.add_page_break()
    doc.add_heading("Soft spectral-outlier sensitivity", level=1)

    target_labels = {
        "fruit_weight_g": "Weight (g)",
        "soluble_solids_pct": "SSC (%)",
        "ph": "pH",
    }
    model_labels = {
        "pls": "Direct PLSR",
        "pls_direct": "Direct PLSR",
        "hierarchical": "Hierarchical PLSR",
        "cnn": "1D-CNN",
        "cnn_texture_aux": "Texture-aux CNN",
        "transformer": "Patch Transformer",
    }

    sensitivity = pd.read_csv(ROOT / "results/sensitivity/soft_spectral_outlier_sensitivity.csv")
    sensitivity = sensitivity[
        [
            "model",
            "target",
            "n_soft_flagged",
            "all_rmse",
            "without_soft_flagged_rmse",
            "delta_rmse_after_exclusion",
            "delta_r2_after_exclusion",
        ]
    ].copy()
    sensitivity["model"] = sensitivity["model"].map(model_labels)
    sensitivity["target"] = sensitivity["target"].map(target_labels)
    for column in sensitivity.columns[3:]:
        sensitivity[column] = sensitivity[column].map(lambda value: f"{value:.4f}")
    sensitivity.columns = ["Model", "Target", "Soft flags", "All RMSE", "RMSE excluding flags", "ΔRMSE", "ΔR²"]
    add_table(
        doc,
        sensitivity,
        "Supplementary Table S1. Sensitivity of zero-shot metrics to 35 soft spectral PCA flags.",
        [1100, 1400, 900, 1300, 1900, 1380, 1380],
        "Negative ΔRMSE denotes lower error after exclusion. Soft flags were retained in every primary analysis.",
    )

    tests = pd.read_csv(ROOT / "results/model_comparison/tables/paired_cultivar_model_tests.csv")
    tests = tests[
        [
            "target",
            "shots",
            "reference_model",
            "competitor_model",
            "reference_mean_cultivar_rmse",
            "competitor_mean_cultivar_rmse",
            "relative_mean_reduction_pct",
            "p_holm_within_target_shots",
        ]
    ].copy()
    tests["target"] = tests["target"].map(target_labels)
    tests["reference_model"] = tests["reference_model"].map(model_labels)
    tests["competitor_model"] = tests["competitor_model"].map(model_labels)
    for column in tests.columns[4:7]:
        tests[column] = tests[column].map(lambda value: f"{value:.4f}")
    tests["p_holm_within_target_shots"] = tests["p_holm_within_target_shots"].map(lambda value: f"{value:.6f}")
    tests.columns = ["Target", "Shots", "Reference", "Competitor", "Reference RMSE", "Competitor RMSE", "Relative reduction (%)", "Holm P"]
    doc.add_page_break()
    doc.add_heading("Paired cultivar-level model tests", level=1)
    add_table(
        doc,
        tests,
        "Supplementary Table S2. Paired cultivar-level Wilcoxon comparisons.",
        [1200, 650, 1250, 1450, 1250, 1350, 1200, 1010],
        "At zero shots direct PLSR is the reference; at 5–50 shots hierarchical PLSR is the reference. Positive relative reduction indicates lower reference-model RMSE. P values are one-sided and Holm adjusted within target and shot count.",
    )

    vip_raw = pd.read_csv(ROOT / "results/spectral_interpretation/top_pls_vip_windows.csv")
    vip = pd.DataFrame(
        {
            "Target": vip_raw["target"].map(target_labels),
            "Window (nm)": vip_raw.apply(lambda row: f"{row.window_low_nm:.0f}–{row.window_high_nm:.0f}", axis=1),
            "Centre (nm)": vip_raw["window_center_nm"].map(lambda value: f"{value:.0f}"),
            "Median VIP": vip_raw["vip_median_across_folds"].map(lambda value: f"{value:.3f}"),
            "VIP IQR": vip_raw.apply(
                lambda row: f"{row.vip_q25_across_folds:.3f}–{row.vip_q75_across_folds:.3f}", axis=1
            ),
            "Folds with mean VIP > 1": vip_raw["fold_fraction_mean_vip_gt_1"].map(lambda value: f"{value:.3f}"),
            "Rank": vip_raw["importance_rank"].map(lambda value: f"{int(value)}"),
        }
    )
    doc.add_heading("Spectral variable importance", level=1)
    add_table(
        doc,
        vip,
        "Supplementary Table S3. Highest mean PLS VIP wavelength windows by target.",
        [1350, 1400, 1200, 1350, 1500, 1700, 860],
        "VIP was computed within each outer held-out-cultivar model. Endpoint regions are susceptible to instrument-edge effects and are not unique molecular assignments.",
    )

    figure = ROOT / "results/spectral_interpretation/figures/figS_pls_vip_stability.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    supplementary_alt_text = (
        "Three wavelength profiles show the median and interquartile range of variable-importance scores "
        "across 16 held-out-cultivar partial least-squares models for fruit weight, soluble solids and pH; "
        "important regions differ by target and include a high endpoint response near 1700 nm."
    )
    shape = p.add_run().add_picture(str(figure), width=Inches(6.35))
    shape._inline.docPr.set("title", "Supplementary Figure S1")
    shape._inline.docPr.set("descr", supplementary_alt_text)
    p = doc.add_paragraph(style="Caption")
    p.paragraph_format.keep_with_next = True
    set_run(
        p.add_run("Supplementary Figure S1. Stability of outer-fold PLS variable-importance-in-projection scores. VIP values are summarized across 16 held-out-cultivar models. Spectral endpoints, especially near 1700 nm, are interpreted as possible instrument-edge sensitivity rather than specific molecular absorptions."),
        size=9,
    )
    p = doc.add_paragraph(style="Table Note")
    set_run(
        p.add_run(f"Alt text: {supplementary_alt_text}"),
        size=8.5,
        italic=True,
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
