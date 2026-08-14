from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_independent_review_report as review


PROJECT_ROOT = Path(__file__).resolve().parents[1]
NAVY = RGBColor(31, 58, 95)
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(35, 47, 62)
MUTED = RGBColor(95, 105, 115)


def configure_narrative_proposal(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def add_editorial_cover(doc: Document) -> None:
    for _ in range(3):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(18)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("FINAL MODEL DEVELOPMENT REPORT")
    review.set_run_font(run, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph(style="Heading 1")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("李子近红外—质构无损预测")
    review.set_run_font(run, size=28, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    run = subtitle.add_run("从 PlumRAC-Net V2.2 到 PLUMRAC-X 的严格跨品种验证")
    review.set_run_font(run, size=15, color=DARK)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(28)
    run = meta.add_run(f"Version 3.0 | {date.today().isoformat()} | Technical review edition")
    review.set_run_font(run, size=10, color=MUTED, italic=True)

    callout = doc.add_paragraph()
    callout.alignment = WD_ALIGN_PARAGRAPH.LEFT
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.left_indent = Inches(0.16)
    callout.paragraph_format.right_indent = Inches(0.16)
    callout.paragraph_format.line_spacing = 1.2
    paragraph_properties = callout._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F6F9")
    paragraph_properties.append(shading)
    run = callout.add_run(
        "核心判断：RD存在小而可信的AI增量，但把模型扩大到336,290参数并未建立九性状普遍优越性；"
        "当前首要瓶颈是跨品种域偏移，而不是隐藏层或激活函数数量不足。"
    )
    review.set_run_font(run, size=10.2, color=DARK, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(26)
    run = paragraph.add_run("Prepared for spectroscopy, statistics, postharvest and AI model review")
    review.set_run_font(run, size=9.5, color=MUTED)
    doc.add_page_break()


def set_report_headers_footers(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run("NIRs_plum | Texture prediction model report")
        review.set_run_font(run, size=8.5, color=MUTED, bold=True)

        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_after = Pt(0)
        review.add_page_field(paragraph)


def main() -> None:
    parser = argparse.ArgumentParser(description="Render the V3 Markdown report as a polished DOCX.")
    parser.add_argument(
        "--markdown",
        type=Path,
        default=PROJECT_ROOT / "reports" / "NIRs_plum_V3_final_model_report_zh.md",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=PROJECT_ROOT / "reports" / "NIRs_plum_V3_final_model_report_zh.docx",
    )
    args = parser.parse_args()
    markdown = args.markdown.resolve().read_text(encoding="utf-8")
    markdown_lines = markdown.splitlines()
    if markdown_lines and markdown_lines[0].startswith("# "):
        markdown_lines = markdown_lines[1:]
        while markdown_lines and not markdown_lines[0].strip():
            markdown_lines = markdown_lines[1:]
        markdown = "\n".join(markdown_lines)

    review.configure_document = configure_narrative_proposal
    review.add_cover = add_editorial_cover
    review.set_headers_footers = set_report_headers_footers
    review.markdown_to_docx(markdown, args.output.resolve())

    document = Document(args.output.resolve())
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "1. 结论先行":
            paragraph._p.addnext(OxmlElement("w:p"))
            break
    document.core_properties.title = "李子近红外—质构无损预测：V3最终模型技术报告"
    document.core_properties.subject = "严格跨品种验证、模型容量消融与部署边界"
    document.core_properties.author = "NIRs_plum project analysis"
    document.core_properties.keywords = "NIR, plum, texture, PLUMRAC-X, PLSR, LOCO"
    document.save(args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
