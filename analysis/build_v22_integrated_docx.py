from __future__ import annotations

import argparse
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT_MD = ROOT / "manuscript" / "manuscript_plumspectra_v22_integrated.md"
SUPPLEMENT_MD = ROOT / "manuscript" / "supplement_plumspectra_v22.md"
MANUSCRIPT_DOCX = ROOT / "manuscript" / "PlumSPECTRA_integrated_manuscript_review.docx"
SUPPLEMENT_DOCX = ROOT / "manuscript" / "PlumSPECTRA_integrated_supplement_review.docx"
SUPPLEMENT_DOCX_ASSETS = ROOT / "results" / "v24_hr_strengthening" / "supplement_docx_assets_w900_c96"
DOCUMENT_RELEASE = "v24"

NAVY = "17243B"
PLUM = "7C3E71"
TEAL = "2E7D75"
GOLD = "C9972B"
MUTED = "617087"
LIGHT_FILL = "F4F6F9"
GRID = "CDD4DE"
WHITE = "FFFFFF"


CAPTIONS = {
    "Figure 1": "Figure 1. Experimental workflow and retained cohort.",
    "Figure 1A artwork (unmerged)": "Figure 1a. Experimental workflow artwork supplied as a separate production file.",
    "Figure 1B quantitative panel (unmerged)": "Figure 1b. Retained-cohort depth supplied as a separate R-rendered production file.",
    "Figure 2": "Figure 2. Integrated phenotype atlas across conventional quality and nine mechanical texture traits.",
    "Figure 3": "Figure 3. PlumSPECTRA architecture and outer-fold performance.",
    "Figure 4": "Figure 4. Complete out-of-fold predictions for all 12 targets.",
    "Figure 5": "Figure 5. Within-cultivar signal and heterogeneous model gain.",
    "Figure 6": "Figure 6. Same-cultivar held-batch audit and deployment boundary.",
}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, *, name: str = "Calibri", size: float | None = None,
                 color: str | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80,
                     start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def prevent_row_split(row) -> None:
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def table_widths(column_count: int) -> list[int]:
    patterns = {
        4: [1200, 2720, 2720, 2720],
        5: [1250, 2190, 1970, 1970, 1980],
        6: [900, 1180, 1580, 1700, 1840, 2160],
        7: [820, 1240, 1250, 1370, 1540, 1350, 1790],
    }
    if column_count in patterns:
        return patterns[column_count]
    base = 9360 // column_count
    widths = [base] * column_count
    widths[-1] += 9360 - sum(widths)
    return widths


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != 9360:
        raise ValueError(f"Table widths must sum to 9360 DXA, got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)

    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), GRID)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_alt_text(inline_shape, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description.split(".")[0])


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_document(doc: Document, running_label: str, right_label: str = "Anonymous review draft") -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        "Heading 1": (16, NAVY, 18, 10),
        "Heading 2": (13, PLUM, 12, 6),
        "Heading 3": (12, NAVY, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.keep_with_next = True

    if "Figure Caption" not in styles:
        caption_style = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption_style = styles["Figure Caption"]
    caption_style.font.name = "Calibri"
    caption_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption_style.font.size = Pt(9)
    caption_style.font.italic = True
    caption_style.font.color.rgb = rgb(MUTED)
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption_style.paragraph_format.space_before = Pt(5)
    caption_style.paragraph_format.space_after = Pt(8)
    caption_style.paragraph_format.keep_with_next = False

    list_tokens = {
        "List Bullet": (0.181, 0.375, -0.194),
        "List Number": (0.25, 0.50, -0.25),
    }
    for list_style_name in ("List Bullet", "List Number"):
        style = styles[list_style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _, text_indent, hanging = list_tokens[list_style_name]
        style.paragraph_format.left_indent = Inches(text_indent)
        style.paragraph_format.first_line_indent = Inches(hanging)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    header = section.header
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(3)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = hp.add_run(running_label)
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = hp.add_run(f"\t{right_label}")
    set_run_font(right, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(3)
    label = fp.add_run("Page ")
    set_run_font(label, size=8.5, color=MUTED)
    add_page_number(fp)

    doc.core_properties.author = "Anonymous"
    doc.core_properties.last_modified_by = "Anonymous"
    doc.core_properties.subject = "PlumSPECTRA integrated NIR phenotyping manuscript"
    doc.core_properties.keywords = "plum, NIR, texture, deep learning, PLSR"


def add_cover(doc: Document, *, title: str, subtitle: str, kicker: str,
              counts: str) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(94)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(kicker.upper())
    set_run_font(run, size=10.5, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10 if subtitle else 26)
    run = p.add_run(title)
    set_run_font(run, size=27, color=NAVY, bold=True)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(26)
        run = p.add_run(subtitle)
        set_run_font(run, size=14, color=PLUM, italic=True)

    if counts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(72)
        run = p.add_run(counts)
        set_run_font(run, size=10.5, color=TEAL, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Anonymous for peer review")
    set_run_font(run, size=11, color=NAVY, bold=True)

    if DOCUMENT_RELEASE != "v29":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        release_date = date.today().strftime("%d %B %Y").lstrip("0")
        run = p.add_run(f"{DOCUMENT_RELEASE.upper()} reviewer manuscript | {release_date}")
        set_run_font(run, size=9.5, color=MUTED)
    p.add_run().add_break(WD_BREAK.PAGE)


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|\[[^\]]+\]\([^\)]+\))")


def add_inline(paragraph, text: str, *, size: float = 11) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, color=NAVY)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=NAVY, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=NAVY, italic=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=max(size - 0.5, 8), color=PLUM)
        else:
            label, target = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", token).groups()
            run = paragraph.add_run(label)
            set_run_font(run, size=size, color=TEAL, bold=True)
            run.font.underline = True
            path_run = paragraph.add_run(f" [{target}]")
            set_run_font(path_run, name="Consolas", size=max(size - 2, 7.5), color=MUTED)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=NAVY)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    column_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=column_count)
    widths = table_widths(column_count)
    for r_index, row_values in enumerate(rows):
        for c_index, value in enumerate(row_values):
            cell = table.cell(r_index, c_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_index < 2 else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            add_inline(paragraph, value, size=8.4 if column_count >= 5 else 9)
            if r_index == 0:
                shade_cell(cell, LIGHT_FILL)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = rgb(NAVY)
    set_repeat_header(table.rows[0])
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        values = [value.strip() for value in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            rows.append(values)
        index += 1
    return rows, index


def prepare_supplement_docx_asset(source: Path) -> Path:
    """Create a browser-oriented figure copy for the <=2 MB supplement file.

    Publication PNG/PDF files are left untouched.  The journal accepts vector
    figure uploads and requires supplementary *files* to be no larger than
    2 MB, so the review/submission DOCX embeds a compact 900-pixel preview
    while the authoritative full-resolution and vector figures remain beside
    it in the release package.
    """
    SUPPLEMENT_DOCX_ASSETS.mkdir(parents=True, exist_ok=True)
    target = SUPPLEMENT_DOCX_ASSETS / source.name
    if target.exists() and target.stat().st_mtime_ns >= source.stat().st_mtime_ns:
        return target
    with Image.open(source) as opened:
        image = opened.convert("RGBA")
        background = Image.new("RGB", image.size, "white")
        background.paste(image, mask=image.getchannel("A"))
        if background.width > 900:
            height = round(background.height * 900 / background.width)
            background = background.resize((900, height), Image.Resampling.LANCZOS)
        compact = background.quantize(colors=96, method=Image.Quantize.MEDIANCUT,
                                      dither=Image.Dither.NONE)
        compact.save(target, format="PNG", optimize=True, compress_level=9)
    return target


def add_figure(doc: Document, source: Path, alt: str, *, supplement: bool,
               legend_lines: list[str] | None = None) -> None:
    if not source.exists():
        raise FileNotFoundError(source)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    # S17 is intentionally tall (two 12-trait training-history grids).  A
    # slightly narrower placement keeps its caption on the same page without
    # changing or downsampling the publication-quality source image.
    compact_main_widths = {
        "Figure 1B quantitative panel (unmerged)": 5.75,
        "Figure 2": 5.70,
        # Figure 4 is portrait-oriented.  At the default 6.45-inch placement
        # its raster height exceeds the 9-inch text block and LibreOffice
        # clips the lower rows.  This changes placement only; the frozen
        # publication figure and its pixels remain untouched.
        "Figure 4": 5.70,
        "Figure 5": 5.80,
    }
    if supplement and alt == "Supplementary Figure S17":
        figure_width = 5.75
    else:
        figure_width = compact_main_widths.get(alt, 6.45)
    embedded_source = prepare_supplement_docx_asset(source) if supplement else source
    shape = paragraph.add_run().add_picture(str(embedded_source), width=Inches(figure_width))
    alt_description = next(
        (line.removeprefix("Alt text:").strip() for line in (legend_lines or [])
         if line.startswith("Alt text:")),
        CAPTIONS.get(alt, alt),
    )
    set_alt_text(shape, alt_description)
    if legend_lines:
        caption_content = [line for line in legend_lines if not line.startswith("Alt text:")]
        caption = doc.add_paragraph(style="Figure Caption")
        caption_size = 8.5 if alt in {"Figure 2", "Figure 6"} else 9
        add_inline(caption, " ".join(caption_content), size=caption_size)
    else:
        caption = doc.add_paragraph(style="Figure Caption")
        caption_text = CAPTIONS.get(alt, alt)
        add_inline(caption, caption_text, size=9)


def resolve_image(markdown_path: Path, target: str) -> Path:
    return (markdown_path.parent / target).resolve()


def render_markdown(doc: Document, markdown_path: Path, *, supplement: bool) -> None:
    lines = markdown_path.read_text(encoding="utf-8-sig").splitlines()
    inline_tables: dict[str, tuple[str, list[list[str]]]] = {}
    inline_figure_legends: dict[str, list[str]] = {}
    if not supplement:
        for scan_index, scan_line in enumerate(lines):
            table_heading = re.match(r"^###\s+(Table\s+\d+)\.\s+(.+)$", scan_line.strip())
            if not table_heading:
                continue
            next_index = scan_index + 1
            while next_index < len(lines) and not lines[next_index].strip():
                next_index += 1
            if next_index < len(lines) and lines[next_index].strip().startswith("|"):
                rows, _ = parse_table(lines, next_index)
                inline_tables[table_heading.group(1)] = (table_heading.group(2), rows)
        for scan_index, scan_line in enumerate(lines):
            legend_heading = re.match(r"^###\s+(Figure\s+\d+)\.\s+(.+)$", scan_line.strip())
            if not legend_heading:
                continue
            legend_lines = [f"{legend_heading.group(1)}. {legend_heading.group(2)}"]
            next_index = scan_index + 1
            while next_index < len(lines):
                candidate = lines[next_index].strip()
                if candidate.startswith("### Figure ") or candidate.startswith("## "):
                    break
                if candidate:
                    legend_lines.append(candidate)
                next_index += 1
            inline_figure_legends[legend_heading.group(1)] = legend_lines
    index = 0
    seen_first_heading = False
    skip_front_metadata = not supplement
    skip_terminal_tables = False
    skip_terminal_figure_legends = False
    inserted_tables: set[str] = set()
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line.startswith("# ") and not seen_first_heading:
            seen_first_heading = True
            index += 1
            continue
        if skip_front_metadata and line.startswith(("**Running title:**", "**Authors:**", "**Affiliations:**", "**Corresponding author:**")):
            index += 1
            continue
        if not supplement and skip_terminal_tables and not line.startswith("## Acknowledgments"):
            index += 1
            continue
        if not supplement and skip_terminal_figure_legends:
            if line == "## References":
                skip_terminal_figure_legends = False
            else:
                index += 1
                continue
        image_match = re.fullmatch(r"!\[([^\]]+)\]\(([^\)]+)\)", line)
        if image_match:
            alt, target = image_match.groups()
            legend_key = "Figure 1" if alt == "Figure 1B quantitative panel (unmerged)" else alt
            legend_lines = None if alt == "Figure 1A artwork (unmerged)" else inline_figure_legends.get(legend_key)
            add_figure(doc, resolve_image(markdown_path, target), alt, supplement=supplement,
                       legend_lines=legend_lines)
            index += 1
            continue
        if line.startswith("|"):
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows)
            continue
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading_match:
            hashes, text = heading_match.groups()
            if not supplement and text == "Tables":
                skip_terminal_tables = True
                index += 1
                continue
            if (not supplement and text == "Figure legends"
                    and DOCUMENT_RELEASE != "v27"):
                skip_terminal_figure_legends = True
                index += 1
                continue
            if not supplement and skip_terminal_tables:
                if text == "Acknowledgments":
                    skip_terminal_tables = False
                else:
                    index += 1
                    continue
            if not supplement and re.match(r"^Table\s+\d+\.", text):
                index += 1
                continue
            level = min(len(hashes) - 1, 3)
            # V26/V27 supply the 24 production figures as separate high-resolution
            # PDF/PNG files.  Keep their legends flowing in the reviewer DOCX
            # instead of forcing one mostly empty page per legend.
            if (supplement and text.startswith("Supplementary Figure")
                    and DOCUMENT_RELEASE not in {"v26", "v27"}):
                doc.add_page_break()
            if not supplement and text == "Tables":
                doc.add_page_break()
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(paragraph, text, size={1: 16, 2: 13, 3: 12}[level])
            index += 1
            continue
        if not supplement and skip_terminal_tables:
            index += 1
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, line[2:], size=10.5)
            index += 1
            continue
        if re.match(r"^\d+\.\s", line):
            paragraph = doc.add_paragraph(style="List Number")
            add_inline(paragraph, re.sub(r"^\d+\.\s*", "", line), size=10.5)
            index += 1
            continue
        paragraph = doc.add_paragraph()
        if supplement and line == "**Anonymous review version**":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "](" in line:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(paragraph, line, size=11)
        if not supplement:
            for table_label, (table_title, table_rows) in inline_tables.items():
                if table_label in inserted_tables or not re.search(rf"\b{re.escape(table_label)}\b", line):
                    continue
                heading = doc.add_paragraph(style="Heading 2")
                add_inline(heading, f"{table_label}. {table_title}", size=13)
                add_markdown_table(doc, table_rows)
                inserted_tables.add(table_label)
        index += 1


def save_doc(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(path)


def build_manuscript() -> None:
    doc = Document()
    if DOCUMENT_RELEASE == "v29":
        configure_document(doc, "PlumSPECTRA | Manuscript", "Anonymous manuscript")
    else:
        configure_document(doc, f"PlumSPECTRA | {DOCUMENT_RELEASE.upper()} manuscript")
    if DOCUMENT_RELEASE in {"v25", "v26", "v27", "v29"}:
        subtitle = "Prediction of multidimensional mechanical texture from intact-fruit spectra"
        kicker = ("Computers and Electronics in Agriculture submission" if DOCUMENT_RELEASE == "v29"
                  else "Systematically audited manuscript revision" if DOCUMENT_RELEASE == "v27"
                  else "Integrated external-review revision" if DOCUMENT_RELEASE == "v26"
                  else "External-review correction manuscript")
        counts = "4,853 plums | 15 cultivars | 12 traits | 58,206 out-of-fold predictions"
    else:
        subtitle = "Nondestructive prediction of conventional quality and nine mechanical texture phenotypes"
        kicker = "Integrated research manuscript"
        counts = "4,839 plums | 15 cultivars | 12 traits | 58,035 out-of-fold predictions"
    cover_title = ("PlumSPECTRA: Cultivar-aware residual learning from near-infrared spectra "
                   "for multidimensional intact-plum texture prediction"
                   if DOCUMENT_RELEASE == "v29" else "PlumSPECTRA")
    if DOCUMENT_RELEASE == "v29":
        subtitle = ""
        counts = ""
    add_cover(
        doc,
        title=cover_title,
        subtitle=subtitle,
        kicker=kicker,
        counts=counts,
    )
    render_markdown(doc, MANUSCRIPT_MD, supplement=False)
    save_doc(doc, MANUSCRIPT_DOCX)


def build_supplement() -> None:
    doc = Document()
    configure_document(doc, "PlumSPECTRA | Supplementary information",
                       "Anonymous manuscript" if DOCUMENT_RELEASE == "v29" else "Anonymous review draft")
    if DOCUMENT_RELEASE in {"v25", "v26", "v27", "v29"}:
        subtitle = ("PlumSPECTRA CEA V29 submission" if DOCUMENT_RELEASE == "v29"
                    else "PlumSPECTRA V27 systematic manuscript revision" if DOCUMENT_RELEASE == "v27"
                    else "PlumSPECTRA V26 integrated revision" if DOCUMENT_RELEASE == "v26"
                    else "PlumSPECTRA V25 external-review correction")
        counts = (("25 supplementary figures | 42 machine-readable tables | 3 frozen manifests")
                  if DOCUMENT_RELEASE == "v29" else
                  ("24 supplementary figures | 40 machine-readable tables | 3 frozen manifests")
                  if DOCUMENT_RELEASE in {"v26", "v27"} else
                  ("25 supplementary figures | 40 machine-readable tables | 3 frozen manifests"))
    else:
        subtitle = "PlumSPECTRA integrated V22/V24 analysis"
        counts = "25 supplementary figures | 32 machine-readable tables | 2 frozen manifests"
    add_cover(
        doc,
        title="Supplementary information",
        subtitle="" if DOCUMENT_RELEASE == "v29" else subtitle,
        kicker="Supplementary material" if DOCUMENT_RELEASE == "v29" else "Figures, tables and reproducibility identifiers",
        counts="" if DOCUMENT_RELEASE == "v29" else counts,
    )
    render_markdown(doc, SUPPLEMENT_MD, supplement=True)
    save_doc(doc, SUPPLEMENT_DOCX)


def main() -> None:
    global MANUSCRIPT_MD, SUPPLEMENT_MD, MANUSCRIPT_DOCX, SUPPLEMENT_DOCX
    global SUPPLEMENT_DOCX_ASSETS, DOCUMENT_RELEASE
    parser = argparse.ArgumentParser()
    parser.add_argument("--document", choices=("manuscript", "supplement", "both"), default="both")
    parser.add_argument("--release", choices=("v24", "v25", "v26", "v27", "v29"), default="v24")
    args = parser.parse_args()
    DOCUMENT_RELEASE = args.release
    if DOCUMENT_RELEASE == "v25":
        MANUSCRIPT_MD = ROOT / "manuscript/manuscript_plumspectra_v25_final.md"
        SUPPLEMENT_MD = ROOT / "manuscript/supplement_plumspectra_v25_final.md"
        MANUSCRIPT_DOCX = ROOT / "manuscript/PlumSPECTRA_V25_manuscript_review.docx"
        SUPPLEMENT_DOCX = ROOT / "manuscript/PlumSPECTRA_V25_supplement_review.docx"
        SUPPLEMENT_DOCX_ASSETS = ROOT / "results/v25_external_review_corrections/supplement_docx_assets"
    elif DOCUMENT_RELEASE == "v26":
        MANUSCRIPT_MD = ROOT / "manuscript/manuscript_plumspectra_v26_integrated.md"
        SUPPLEMENT_MD = ROOT / "manuscript/supplement_plumspectra_v26_integrated.md"
        MANUSCRIPT_DOCX = ROOT / "manuscript/PlumSPECTRA_V26_integrated_manuscript_review.docx"
        SUPPLEMENT_DOCX = ROOT / "manuscript/PlumSPECTRA_V26_integrated_supplement_review.docx"
        SUPPLEMENT_DOCX_ASSETS = ROOT / "results/v26_claudecode_integration/supplement_docx_assets"
    elif DOCUMENT_RELEASE == "v27":
        MANUSCRIPT_MD = ROOT / "manuscript/manuscript_plumspectra_v27.md"
        SUPPLEMENT_MD = ROOT / "manuscript/supplement_plumspectra_v27.md"
        MANUSCRIPT_DOCX = ROOT / "manuscript/PlumSPECTRA_V27_manuscript_review.docx"
        SUPPLEMENT_DOCX = ROOT / "manuscript/PlumSPECTRA_V27_supplement_review.docx"
        SUPPLEMENT_DOCX_ASSETS = ROOT / "results/v26_claudecode_integration/supplement_docx_assets"
    elif DOCUMENT_RELEASE == "v29":
        MANUSCRIPT_MD = ROOT / "manuscript/manuscript_plumspectra_cea_v29.md"
        SUPPLEMENT_MD = ROOT / "manuscript/supplement_plumspectra_cea_v29.md"
        MANUSCRIPT_DOCX = ROOT / "manuscript/cea_v29_submission/PlumSPECTRA_CEA_manuscript_revised.docx"
        SUPPLEMENT_DOCX = ROOT / "manuscript/cea_v29_submission/PlumSPECTRA_CEA_supplement_revised.docx"
        SUPPLEMENT_DOCX_ASSETS = ROOT / "results/v29_cea_submission/supplement_docx_assets"
    if args.document in ("manuscript", "both"):
        build_manuscript()
    if args.document in ("supplement", "both"):
        build_supplement()


if __name__ == "__main__":
    main()
