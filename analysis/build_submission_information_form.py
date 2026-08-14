from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "manuscript" / "Horticulture_Research_submission_information_form.docx"

INK = "17233B"
PLUM = "7B2D5F"
TEAL = "238C86"
PALE_PLUM = "F5EAF1"
PALE_GOLD = "FFF4D6"
PALE_TEAL = "E7F4F2"
PALE_GREY = "F2F4F7"
WHITE = "FFFFFF"
MUTED = "65738A"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn("w:" + key), str(value))


def set_cell_margins(cell, top=85, start=95, bottom=85, end=95) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def run_text(paragraph, text: str, *, bold=False, color=INK, size=9.0, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def label_paragraph(cell, text: str, color=INK, size=8.6, bold=False):
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run_text(paragraph, text, bold=bold, color=color, size=size)
    return paragraph


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    for style_name, size, color in [
        ("Title", 21, INK), ("Heading 1", 14, PLUM), ("Heading 2", 10.5, INK)
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_text(p, "PLUMSPECTRA  /  SUBMISSION METADATA", bold=True, color=PLUM, size=8)
    footer = section.footer
    p = footer.paragraphs[0]
    run_text(p, "CONFIDENTIAL AUTHOR WORKSHEET  •  complete highlighted cells before submission", color=MUTED, size=7.5)
    add_page_number(footer.add_paragraph())
    return doc


def add_status_legend(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for cell, label, fill, color in zip(
        table.rows[0].cells,
        ["CONFIRMED FROM ARCHIVE", "AUTHOR MUST COMPLETE", "OPTIONAL / IF APPLICABLE"],
        [PALE_TEAL, PALE_GOLD, PALE_GREY],
        [TEAL, "8B5E00", MUTED],
    ):
        set_cell_width(cell, 5.8)
        shade(cell, fill)
        set_cell_border(cell, bottom={"val": "single", "sz": 8, "color": WHITE})
        label_paragraph(cell, label, color=color, size=7.8, bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section_intro(doc: Document, heading: str, instruction: str) -> None:
    doc.add_heading(heading, level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run_text(p, instruction, color=MUTED, size=8.5, italic=True)


def add_form_table(doc: Document, rows: list[tuple[str, str, str]], widths=(5.0, 11.6)) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, width, text in zip(header.cells, widths, ["ITEM", "RESPONSE / VERIFIED VALUE"]):
        set_cell_width(cell, width)
        shade(cell, PLUM)
        label_paragraph(cell, text, color=WHITE, size=8.2, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for label, value, status in rows:
        row = table.add_row()
        prevent_row_split(row)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.72)
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(
                cell,
                bottom={"val": "single", "sz": 5, "color": "D7DDE5"},
                left={"val": "single", "sz": 3, "color": "E7EAF0"},
                right={"val": "single", "sz": 3, "color": "E7EAF0"},
            )
        shade(row.cells[0], PALE_GREY)
        shade(row.cells[1], PALE_TEAL if status == "confirmed" else PALE_GOLD if status == "required" else WHITE)
        label_paragraph(row.cells[0], label, bold=True, size=8.4)
        display = value if value else "Click here and enter author-verified information."
        label_paragraph(
            row.cells[1], display,
            color=INK if status == "confirmed" else "8B5E00" if status == "required" else MUTED,
            size=8.6,
        )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_checklist(doc: Document, heading: str, items: list[str]) -> None:
    doc.add_heading(heading, level=2)
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.2)
        p.paragraph_format.space_after = Pt(2.5)
        run_text(p, item, size=8.6)


def build() -> None:
    doc = setup_document()
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_text(title, "Horticulture Research\nsubmission information form", bold=True, color=INK, size=21)
    subtitle = doc.add_paragraph()
    run_text(subtitle, "PlumSPECTRA • author completion worksheet • version 2026-08-08", bold=True, color=PLUM, size=10)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run_text(
        p,
        "Purpose. This form separates facts recovered from the source archive from information that only the research team can verify. Do not replace highlighted blanks with guesses. Return the completed file before final journal submission.",
        color=MUTED,
        size=9,
    )
    add_status_legend(doc)

    add_section_intro(doc, "1  Manuscript identity and governance", "Confirm the final submission identity. Every author should approve the order, contributions and declarations.")
    add_form_table(doc, [
        ("Working title", "PlumSPECTRA unifies nondestructive prediction of conventional quality and nine mechanical texture phenotypes across 4,839 plums", "confirmed"),
        ("Article type", "Research Article", "required"),
        ("Authors and order", "", "required"),
        ("Affiliations", "", "required"),
        ("Corresponding author", "Name, full postal address and email:", "required"),
        ("ORCID identifiers", "", "required"),
        ("CRediT contributions", "", "required"),
        ("Funding", "Agency names and grant numbers; state ‘None’ only if correct:", "required"),
        ("Acknowledgements", "", "required"),
        ("Conflict of interest", "", "required"),
        ("Data redistribution rights", "☐ Confirmed  ☐ Restricted  ☐ Needs institutional review", "required"),
    ])

    doc.add_page_break()
    add_section_intro(doc, "2  Biological material and sampling", "These details are not recoverable from the current archive and are essential for horticultural reproducibility.")
    add_form_table(doc, [
        ("Species / taxon", "Scientific name and authority; clarify Japanese plum, hybrids or other material:", "required"),
        ("Material status", "Identify registered cultivars versus breeding selections; confirm LA191 as the publication label for archived A181.", "required"),
        ("Cultivar provenance", "Source nursery, breeding programme or germplasm collection for each material, if known:", "required"),
        ("Orchard / station", "Name, municipality, province/state and country:", "required"),
        ("Trees and design", "Number of trees per cultivar, block design and fruit per tree:", "required"),
        ("Harvest timing", "Year(s), exact or approximate dates, and time of day if relevant:", "required"),
        ("Maturity criteria", "Colour, firmness, days after bloom, commercial maturity or other picking rule:", "required"),
        ("Sampling strategy", "Canopy position, randomisation, size/defect selection and biological replication:", "required"),
        ("Postharvest interval", "Time from harvest to each measurement stage:", "required"),
        ("Storage / transport", "Temperature, humidity, packaging and equilibration:", "required"),
        ("Batch definition", "The archive contains 19 source batches. State whether codes denote dates, lots, trees, operators or acquisition sessions.", "required"),
        ("Operator changes", "Who measured which batches, and when personnel or protocol changed:", "required"),
    ])

    doc.add_page_break()
    add_section_intro(doc, "3  NIR acquisition", "Archive-derived acquisition fields are prefilled. Supply manufacturer, geometry and referencing details from laboratory records if available.")
    add_form_table(doc, [
        ("Archived serial number", "5490277", "confirmed"),
        ("Spectral configuration", "900–1700 nm configured range; 228 digital channels; analysed absorbance grid approximately 901–1701 nm.", "confirmed"),
        ("Exposure and repeats", "5.080 ms exposure; 6 repeats recorded in the source export.", "confirmed"),
        ("Detector metadata", "Representative export: system 42.3 °C, detector 42.34 °C, humidity 19.8%; do not treat these as fruit temperature.", "confirmed"),
        ("Manufacturer / model", "Spectrometer corresponding to serial 5490277:", "required"),
        ("Optical geometry", "Reflectance, interactance or transmittance; contact/non-contact; probe/accessory:", "required"),
        ("Fruit scan location", "Anatomical side, orientation, number of positions and whether peel remained intact:", "required"),
        ("Replicate handling", "Confirm whether the six archived repeats were averaged by the instrument and whether additional scans were acquired.", "required"),
        ("Reference schedule", "White/dark reference material, frequency, warm-up and recalibration schedule:", "required"),
        ("Codes c and t", "Define archived scan suffixes. The analysis uses c as the primary scan but assigns no unverified geometry.", "required"),
        ("Sample conditions", "Fruit surface preparation, ambient and fruit temperature, lighting and operator protocol:", "required"),
        ("Timestamp discrepancy", "Explain any difference between embedded host timestamps and workbook/file dates before reporting acquisition dates.", "required"),
    ])

    doc.add_page_break()
    add_section_intro(doc, "4  Texture acquisition and reference phenotypes", "Confirm settings that determine the physical meaning of each mechanical endpoint and conventional reference measurement.")
    add_form_table(doc, [
        ("Software family", "Stable Micro Systems Exponent archive output", "confirmed"),
        ("Probe", "P/2, 2-mm-diameter cylindrical stainless-steel probe", "confirmed"),
        ("Sampling and load cell", "Nominal 400 Hz; archived load-cell capacity 30,000 g", "confirmed"),
        ("Force unit evidence", "Raw force ×1000 reproduced Force (g) previews across 19/19 representative source batches; results reported as gf.", "confirmed"),
        ("Position-unit boundary", "Calibrated distance conversion is unavailable; displacement, stiffness and work remain in raw position units.", "confirmed"),
        ("Texture analyser", "Manufacturer/model and serial number:", "required"),
        ("Exponent version", "Software version and method/template name:", "required"),
        ("Motion settings", "Pre-test, test and post-test speeds; trigger force; target distance/depth; return settings:", "required"),
        ("Fruit positioning", "Orientation, support fixture and exact anatomical positions for replicates 01 and 02:", "required"),
        ("Instrument calibration", "Load-cell and distance calibration date/procedure:", "required"),
        ("Balance", "Manufacturer/model, resolution, calibration and weighing protocol:", "required"),
        ("Soluble solids", "Refractometer manufacturer/model, calibration, juice preparation and confirm whether % means °Brix:", "required"),
        ("pH", "Meter/model, electrode, pH 4/7/10 buffers, temperature correction, juice preparation and replicates:", "required"),
        ("Target confirmation", "☐ Confirmed: the acidity-related target is pH, not titratable acidity.", "required"),
    ])

    doc.add_page_break()
    add_section_intro(doc, "5  Data, code and submission declarations", "Complete permanent repository and journal-facing declarations only after the release package is frozen.")
    add_form_table(doc, [
        ("Dataset release", "Repository, version, DOI and private reviewer link:", "required"),
        ("Code release", "Repository, tagged release, archive DOI and private reviewer link:", "required"),
        ("Licences", "Dataset licence; code licence; restrictions on source-derived ARC/DAT files:", "required"),
        ("Availability wording", "State exactly which raw, cleaned and derived data can be redistributed and which require controlled access.", "required"),
        ("Ethics / permits", "Plant-material permits, access permissions or ‘Not applicable’, as verified by the authors:", "required"),
        ("AI disclosure", "Confirm journal-required disclosure for AI-assisted language editing/figure workflow, if applicable:", "required"),
        ("Preprint", "Server and DOI, or ‘None’:", "optional"),
        ("Suggested reviewers", "Names, institutions, emails, expertise and conflict check:", "optional"),
        ("Opposed reviewers", "Names and concise conflict rationale, if applicable:", "optional"),
    ])
    add_checklist(doc, "Final author verification", [
        "All yellow cells are completed or explicitly marked ‘unknown after author verification’. Do not invent missing protocol details.",
        "Cultivar names, breeding-selection labels and LA191/A181 mapping are approved by the germplasm owner.",
        "The public data package contains no private Windows paths or personally identifying source filenames.",
        "Repository files reproduce the frozen outer-fold manifests, predictions, figures and tables reported in the manuscript.",
        "All authors approve the submitted manuscript, contribution statement, funding, data rights and conflicts.",
    ])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run_text(p, "Completed by: ______________________________", bold=True, size=9)
    run_text(p, "     Date: __________________     Version returned: __________________", size=9)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
