#!/usr/bin/env python3
"""Audit the integrated V22/V24 PlumSPECTRA Horticulture Research package.

The audit is intentionally independent of model training. It verifies that the
frozen prediction evidence, quantitative figures, supplementary material, and
review documents agree on the main cohort and claim boundaries.
"""

from __future__ import annotations

import hashlib
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZipFile

import pandas as pd
from docx import Document
from lxml import etree
from PIL import Image
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results" / "v22_integrated"
FIGURE_DATA = RESULTS / "figure_data"
FIGURES = RESULTS / "figures_r"
SUPPLEMENT = RESULTS / "supplement"
MANUSCRIPT = ROOT / "manuscript"
AUDIT_DIR = ROOT / "review_package" / "final_release_audit"
AUDIT_JSON = AUDIT_DIR / "v24_hr_release_audit.json"
V24 = ROOT / "results" / "v24_hr_strengthening"
HR_PACKAGE = MANUSCRIPT / "HR_submission_package_v24"

TEXTURE_HASH = "f44902e12579b033c354c41a0c00f681801218b3bca5ddc52d7e9cee7dba4105"
QUALITY_HASH = "4389b2201e1c79f2b4c7b28c36829fcb1093ece350e9de5cf5a95ecd04c25379"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


class Audit:
    def __init__(self) -> None:
        self.checks: list[dict[str, object]] = []

    def check(self, name: str, passed: bool, observed: object, expected: object) -> None:
        self.checks.append(
            {
                "name": name,
                "status": "PASS" if passed else "FAIL",
                "observed": observed,
                "expected": expected,
            }
        )

    def require_file(self, path: Path, name: str | None = None) -> None:
        present = path.is_file() and path.stat().st_size > 0
        self.check(name or f"file:{path.relative_to(ROOT)}", present, present, True)

    def finish(self) -> dict[str, object]:
        failures = [entry for entry in self.checks if entry["status"] == "FAIL"]
        return {
            "release": "PlumSPECTRA V22/V24 Horticulture Research candidate",
            "generated_utc": datetime.now(timezone.utc).isoformat(),
            "status": "PASS" if not failures else "FAIL",
            "summary": {
                "checks": len(self.checks),
                "passed": len(self.checks) - len(failures),
                "failed": len(failures),
            },
            "checks": self.checks,
        }


def docx_xml(path: Path) -> etree._Element:
    with ZipFile(path) as archive:
        return etree.fromstring(archive.read("word/document.xml"))


def docx_text(doc: Document) -> str:
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def audit_docx(audit: Audit, path: Path, label: str, expected_tables: int, expected_images: int) -> None:
    audit.require_file(path, f"{label}:DOCX exists")
    doc = Document(path)
    section = doc.sections[0]
    geometry = {
        "page_width_twips": section.page_width.twips,
        "page_height_twips": section.page_height.twips,
        "top_margin_twips": section.top_margin.twips,
        "bottom_margin_twips": section.bottom_margin.twips,
        "left_margin_twips": section.left_margin.twips,
        "right_margin_twips": section.right_margin.twips,
        "header_distance_twips": section.header_distance.twips,
        "footer_distance_twips": section.footer_distance.twips,
    }
    expected_geometry = {
        "page_width_twips": 12240,
        "page_height_twips": 15840,
        "top_margin_twips": 1440,
        "bottom_margin_twips": 1440,
        "left_margin_twips": 1440,
        "right_margin_twips": 1440,
        "header_distance_twips": 708,
        "footer_distance_twips": 708,
    }
    audit.check(f"{label}:US Letter geometry", geometry == expected_geometry, geometry, expected_geometry)
    audit.check(f"{label}:table count", len(doc.tables) == expected_tables, len(doc.tables), expected_tables)
    audit.check(f"{label}:image count", len(doc.inline_shapes) == expected_images, len(doc.inline_shapes), expected_images)

    root = docx_xml(path)
    exact_row_heights = root.xpath(".//w:trHeight[@w:hRule='exact']", namespaces=NS)
    audit.check(f"{label}:no exact row heights", len(exact_row_heights) == 0, len(exact_row_heights), 0)

    if expected_tables:
        expected_width = 9360
        widths = [int(value) for value in root.xpath(".//w:tblW/@w:w", namespaces=NS)]
        indents = [int(value) for value in root.xpath(".//w:tblInd/@w:w", namespaces=NS)]
        audit.check(
            f"{label}:table widths fit text area",
            bool(widths) and all(value == expected_width for value in widths),
            widths,
            [expected_width] * expected_tables,
        )
        audit.check(
            f"{label}:table indents are consistent",
            bool(indents) and all(value == 120 for value in indents),
            indents,
            [120] * expected_tables,
        )


def main() -> int:
    audit = Audit()

    main_docx = MANUSCRIPT / "PlumSPECTRA_integrated_manuscript_review.docx"
    supp_docx = MANUSCRIPT / "PlumSPECTRA_integrated_supplement_review.docx"
    main_pdf = MANUSCRIPT / "PlumSPECTRA_integrated_manuscript_review.pdf"
    supp_pdf = MANUSCRIPT / "PlumSPECTRA_integrated_supplement_review.pdf"
    main_md = MANUSCRIPT / "manuscript_plumspectra_v22_integrated.md"
    supp_md = MANUSCRIPT / "supplement_plumspectra_v22.md"

    audit_docx(audit, main_docx, "main", expected_tables=3, expected_images=7)
    audit_docx(audit, supp_docx, "supplement", expected_tables=0, expected_images=25)

    for path, pages, label in [(main_pdf, 29, "main"), (supp_pdf, 33, "supplement")]:
        audit.require_file(path, f"{label}:PDF exists")
        page_count = len(PdfReader(path).pages)
        audit.check(f"{label}:PDF page count", page_count == pages, page_count, pages)

    for path, label in [(main_md, "main markdown"), (supp_md, "supplement markdown")]:
        audit.require_file(path, f"{label} exists")

    main_text = main_md.read_text(encoding="utf-8")
    supp_text = supp_md.read_text(encoding="utf-8")
    doc_text = docx_text(Document(main_docx))
    abstract_match = re.search(r"^## Abstract\s+(.*?)(?=^## )", main_text, flags=re.MULTILINE | re.DOTALL)
    abstract_body = re.split(r"^\*{0,2}Keywords", abstract_match.group(1), maxsplit=1,
                             flags=re.MULTILINE | re.IGNORECASE)[0] if abstract_match else ""
    abstract_words = len(re.findall(r"\b[\w'-]+\b", abstract_body))
    references_match = re.search(r"^## References\s+(.*?)(?=^## |\Z)", main_text, flags=re.MULTILINE | re.DOTALL)
    reference_count = len(re.findall(r"^\d+\.\s", references_match.group(1), flags=re.MULTILINE)) if references_match else 0
    body_without_backmatter = re.split(r"^## References\s*$", main_text, maxsplit=1, flags=re.MULTILINE)[0]
    body_words = len(re.findall(r"\b[\w'-]+\b", re.sub(r"!\[[^]]*\]\([^)]*\)", "", body_without_backmatter)))
    main_figure_count = len(re.findall(r"^### Figure [1-6]\. ", main_text, flags=re.MULTILINE))
    main_table_count = len(re.findall(r"^### Table [1-3]\. ", main_text, flags=re.MULTILINE))
    alt_text_count = len(re.findall(r"^\*{0,2}Alt text", main_text, flags=re.MULTILINE))
    audit.check("HR abstract word limit", 0 < abstract_words <= 250, abstract_words, "<= 250")
    audit.check("HR article word limit", body_words <= 6000, body_words, "<= 6000")
    audit.check("HR reference limit", reference_count <= 50, reference_count, "<= 50")
    audit.check("HR main figure limit", main_figure_count == 6, main_figure_count, 6)
    audit.check("HR main table limit", main_table_count == 3, main_table_count, 3)
    audit.check("HR figure alt-text coverage", alt_text_count >= 6, alt_text_count, ">= 6")
    audit.check("required Supplementary information section", "## Supplementary information" in main_text,
                "## Supplementary information" in main_text, True)
    required_terms = [
        "PlumSPECTRA",
        "Plum Spectral Phenotyping Ensemble with Cultivar-aware Trait Residual Adaptation",
        "4,839",
        "4,828",
        "58,035",
        "fruit weight",
        "soluble solids content",
        "pH",
        "nine texture",
    ]
    for term in required_terms:
        audit.check(
            f"main manuscript term:{term}",
            term.lower() in main_text.lower() and term.lower() in doc_text.lower(),
            {"markdown": term.lower() in main_text.lower(), "docx": term.lower() in doc_text.lower()},
            {"markdown": True, "docx": True},
        )
    pH_boundary = bool(re.search(r"pH[^\n]{0,180}not titratable acidity", main_text, flags=re.IGNORECASE))
    audit.check("pH is not mislabelled as titratable acidity", pH_boundary, pH_boundary, True)
    mapping_correct = "A181" in main_text and "LA191" in main_text
    audit.check("A181 to LA191 correction retained", mapping_correct, mapping_correct, True)

    predictions = pd.read_csv(FIGURE_DATA / "predictions_all12.csv")
    pooled = pd.read_csv(FIGURE_DATA / "pooled_metrics.csv")
    comparisons = pd.read_csv(FIGURE_DATA / "final_model_cluster_comparisons.csv")
    counts = pd.read_csv(FIGURE_DATA / "cohort_counts.csv")
    audit.check("integrated prediction rows", len(predictions) == 58035, len(predictions), 58035)
    audit.check("integrated traits", predictions["trait"].nunique() == 12, predictions["trait"].nunique(), 12)
    family_rows = predictions.groupby("family").size().to_dict()
    expected_family_rows = {"Conventional quality": 14484, "Mechanical texture": 43551}
    audit.check(
        "prediction rows by family",
        family_rows == expected_family_rows,
        family_rows,
        expected_family_rows,
    )
    audit.check("cultivar count", counts["cultivar_ascii"].nunique() == 15, counts["cultivar_ascii"].nunique(), 15)
    texture_fruit = int(family_rows["Mechanical texture"] / 9)
    quality_fruit = int(family_rows["Conventional quality"] / 3)
    audit.check("texture fruit count", texture_fruit == 4839, texture_fruit, 4839)
    audit.check("quality fruit count", quality_fruit == 4828, quality_fruit, 4828)
    audit.check("final comparison count", len(comparisons) == 36, len(comparisons), 36)
    supported = int((comparisons["claim_status"].str.lower() == "supported_outperformance").sum())
    audit.check("CI-supported comparisons", supported == 35, supported, 35)
    audit.check("all traits have final metrics", pooled.loc[pooled["is_final"].astype(bool), "trait"].nunique() == 12,
                pooled.loc[pooled["is_final"].astype(bool), "trait"].nunique(), 12)

    quality_manifest = RESULTS / "splits" / "v22_quality_fivefold_manifest.csv"
    audit.check("quality manifest SHA-256", sha256(quality_manifest) == QUALITY_HASH, sha256(quality_manifest), QUALITY_HASH)
    hash_terms_present = TEXTURE_HASH in supp_text and QUALITY_HASH in supp_text
    audit.check("frozen manifest hashes reported in supplement", hash_terms_present, hash_terms_present, True)

    main_figure_stems = [
        "fig01b_cohort_depth",
        "fig02_integrated_phenotype_atlas",
        "fig03_plumspectra_architecture_performance",
        "fig04_all12_observed_predicted",
        "fig05_within_cultivar_heterogeneity",
        "fig06_crossbatch_boundary",
    ]
    for stem in main_figure_stems:
        for suffix in [".png", ".pdf"]:
            audit.require_file(FIGURES / f"{stem}{suffix}")
    audit.require_file(FIGURES / "fig03a_plumspectra_architecture.png", "standalone architecture PNG")
    audit.require_file(FIGURES / "fig03a_plumspectra_architecture.pdf", "standalone architecture PDF")
    audit.require_file(FIGURES / "fig03a_plumspectra_architecture.svg", "editable architecture SVG")
    audit.require_file(FIGURES / "figure_manifest.csv", "main figure manifest")

    audit.require_file(
        RESULTS / "imagegen" / "IMAGEGEN_PROMPT_MANIFEST.md",
        "ImageGen prompt manifest",
    )
    workflow_source_candidates = sorted((RESULTS / "imagegen" / "fig01a_candidates").glob("*.png"))
    graphical_source_candidates = sorted((RESULTS / "imagegen" / "graphical_abstract_candidates").glob("*.png"))
    audit.check("ImageGen workflow source candidate count", len(workflow_source_candidates) == 5,
                len(workflow_source_candidates), 5)
    audit.check("ImageGen graphical source candidate count", len(graphical_source_candidates) == 5,
                len(graphical_source_candidates), 5)
    figure_manifest = pd.read_csv(FIGURES / "figure_manifest.csv")
    imagegen_rows = figure_manifest[figure_manifest["source"].eq("ImageGen (unmerged)")]
    audit.check("ImageGen and R panels remain unmerged", len(imagegen_rows) == 10,
                len(imagegen_rows), 10)

    supplement_pngs = sorted((SUPPLEMENT / "figures").glob("figS*.png"))
    supplement_pdfs = sorted((SUPPLEMENT / "figures").glob("figS*.pdf"))
    supplement_tables = sorted((SUPPLEMENT / "tables").glob("*.csv"))
    audit.check("supplement figure PNG count", len(supplement_pngs) == 25, len(supplement_pngs), 25)
    audit.check("supplement figure PDF count", len(supplement_pdfs) == 25, len(supplement_pdfs), 25)
    audit.check("supplement table count", len(supplement_tables) == 32, len(supplement_tables), 32)
    two_mb = 2 * 1024 * 1024
    audit.check("supplement DOCX is within 2 MB", supp_docx.stat().st_size <= two_mb,
                supp_docx.stat().st_size, f"<= {two_mb} bytes")
    audit.check("supplement PDF is within 2 MB", supp_pdf.stat().st_size <= two_mb,
                supp_pdf.stat().st_size, f"<= {two_mb} bytes")

    r_script = ROOT / "src" / "render_v22_integrated_figures.R"
    audit.require_file(r_script, "R quantitative figure script")
    r_text = r_script.read_text(encoding="utf-8")
    audit.check("ggplot2 used", "library(ggplot2)" in r_text, "library(ggplot2)" in r_text, True)
    native_correlation_heatmap = (
        "corr_panel <- ggplot" in r_text
        and "geom_tile" in r_text
        and "coord_fixed" in r_text
        and "library(ComplexHeatmap)" not in r_text
    )
    audit.check("correlation matrix uses a native square ggplot heatmap",
                native_correlation_heatmap, native_correlation_heatmap, True)
    classic2_only = (
        "theme_classic2 <- function" in r_text
        and "theme_classic(base_size" in r_text
        and "theme_minimal" not in r_text
        and "theme_journal" not in r_text
    )
    audit.check("all R figures use the classic2 theme system", classic2_only, classic2_only, True)
    arial_registered = (
        'font_family <- "Arial"' in r_text
        and "windowsFonts(Arial" in r_text
        and "postscriptFonts(Arial" in r_text
        and "pdfFonts(Arial" in r_text
    )
    audit.check("Arial is registered across Windows and vector devices",
                arial_registered, arial_registered, True)
    ten_half_width = (
        "standard_width <- 10.5" in r_text
        and "standard_dpi <- 450" in r_text
        and "publication_dimensions" in r_text
    )
    audit.check("R export policy fixes non-Figure-1B width at 10.5 inches",
                ten_half_width, ten_half_width, True)
    open_categorical_axes = (
        "theme_categorical_x <- function" in r_text
        and "theme_categorical_y <- function" in r_text
        and "axis.line.x = element_blank()" in r_text
        and "axis.line.y = element_blank()" in r_text
    )
    audit.check("categorical axes use an open-frame design",
                open_categorical_axes, open_categorical_axes, True)
    zero_expansion_policy = (
        "zero_lower_expand <- expansion(mult = c(0, 0.04))" in r_text
        and r_text.count("expand = zero_lower_expand") >= 8
    )
    audit.check("zero-based bars and lollipops have no lower-side expansion",
                zero_expansion_policy, zero_expansion_policy, True)
    no_panel_titles = (
        not re.search(r"labs\([^)]*\b(?:title|subtitle)\s*=", r_text, flags=re.DOTALL)
        and "ggtitle(" not in r_text
        and "plot_annotation(" not in r_text
    )
    audit.check("R panels contain tags but no plot titles", no_panel_titles, no_panel_titles, True)
    count_block = r_text[r_text.find("count_plot <-"):r_text.find("# ---- Figure 2")]
    count_has_no_labels = "geom_col" in count_block and "geom_text" not in count_block and "geom_label" not in count_block
    audit.check("Figure 1B has no bar-top numeric labels", count_has_no_labels, count_has_no_labels, True)

    fig2_block = r_text[r_text.find("# ---- Figure 2"):r_text.find("# ---- Figure 3")]
    fig2_distribution_design = (
        fig2_block.count("geom_boxplot") >= 2
        and fig2_block.count("position_jitter") >= 2
    )
    audit.check("Figure 2A/B use boxplots plus fruit-level jitter", fig2_distribution_design,
                fig2_distribution_design, True)
    pca_centroid_design = (
        "stat_ellipse" not in fig2_block
        and "pca_centroids" in fig2_block
        and "geom_text_repel(data = pca_centroids" in fig2_block
        and "geom_point(data = pca_centroids" in fig2_block
    )
    audit.check("PCA uses labelled cultivar centroids without confidence ellipses",
                pca_centroid_design, pca_centroid_design, True)

    fig3_block = r_text[r_text.find("# ---- Figure 3"):r_text.find("# ---- Figure 4")]
    fig3_uncertainty = (
        fig3_block.count("geom_errorbar") >= 2
        and "comparison_evidence" in fig3_block
        and "95% CI excludes zero" in fig3_block
        and "95% CI includes zero" in fig3_block
    )
    audit.check("Figure 3 shows uncertainty or explicit cluster-evidence status", fig3_uncertainty,
                fig3_uncertainty, True)
    r2_is_interval_plot = "r2_plot <-" in fig3_block and "geom_point" in fig3_block and "geom_errorbar" in fig3_block
    audit.check("Figure 3 R-squared uses point-interval chart", r2_is_interval_plot, r2_is_interval_plot, True)
    redundant_fig3_zero_lines_absent = (
        "geom_hline(yintercept = 0" not in fig3_block
        and "geom_vline(xintercept = 0" not in fig3_block
    )
    audit.check("Figure 3 omits redundant zero frame lines",
                redundant_fig3_zero_lines_absent, redundant_fig3_zero_lines_absent, True)

    non_fig1b_pngs = [
        *(FIGURES / f"{stem}.png" for stem in main_figure_stems if stem != "fig01b_cohort_depth"),
        FIGURES / "fig03a_plumspectra_architecture.png",
        *supplement_pngs,
    ]
    raster_specs: dict[str, dict[str, object]] = {}
    raster_ok = True
    for path in [FIGURES / "fig01b_cohort_depth.png", *non_fig1b_pngs]:
        with Image.open(path) as image:
            dpi = image.info.get("dpi", (None, None))
            observed = {
                "width_px": image.width,
                "height_px": image.height,
                "dpi_x": None if dpi[0] is None else round(float(dpi[0]), 2),
                "dpi_y": None if dpi[1] is None else round(float(dpi[1]), 2),
            }
        raster_specs[str(path.relative_to(ROOT))] = observed
        if path.name == "fig01b_cohort_depth.png":
            expected_width, expected_dpi = 3250, 500
        else:
            expected_width, expected_dpi = 4725, 450
        dpi_ok = observed["dpi_x"] is not None and abs(float(observed["dpi_x"]) - expected_dpi) < 0.5
        raster_ok = raster_ok and observed["width_px"] == expected_width and dpi_ok
    audit.check("all R PNG widths and resolutions follow the publication policy",
                raster_ok, raster_specs,
                {"Figure 1B": "3250 px at 500 dpi", "all other R PNGs": "4725 px at 450 dpi"})

    vector_pdfs = [
        *(FIGURES / f"{stem}.pdf" for stem in main_figure_stems),
        FIGURES / "fig03a_plumspectra_architecture.pdf",
        *supplement_pdfs,
    ]
    pdf_fonts: dict[str, list[str]] = {}
    arial_pdf_ok = True
    for path in vector_pdfs:
        names: set[str] = set()
        for page in PdfReader(path).pages:
            resources = page.get("/Resources", {})
            for font in resources.get("/Font", {}).values():
                names.add(str(font.get_object().get("/BaseFont")))
        pdf_fonts[str(path.relative_to(ROOT))] = sorted(names)
        arial_pdf_ok = arial_pdf_ok and bool(names) and all(
            "Arial" in name or "SymbolMT" in name for name in names
        )
    audit.check("all vector figure PDFs embed Arial, with SymbolMT allowed only for mathematical glyphs",
                arial_pdf_ok, pdf_fonts, "Every BaseFont contains 'Arial' or is SymbolMT for mathematical notation")

    supplement_loop = r_text[r_text.find("for (i in seq_along(trait_levels))"):r_text.find("# ---- Figure 5")]
    supplement_facet_logic = "facet_wrap(~cultivar_code" in supplement_loop and "geom_point" in supplement_loop
    audit.check("supplement uses cultivar-resolved facet plots", supplement_facet_logic,
                supplement_facet_logic, True)
    model_plot_block = r_text[r_text.find("fig03_plumspectra_architecture_performance") - 18000:]
    r2_tile_absent = not bool(re.search(r"aes\([^)]*r2[^)]*\)[^\n]{0,300}geom_tile", model_plot_block, flags=re.IGNORECASE))
    audit.check("no R-squared performance heatmap", r2_tile_absent, r2_tile_absent, True)
    cultivar_trait_grid = predictions.groupby("trait")["cultivar_code"].nunique().to_dict()
    all_traits_have_15_cultivars = len(cultivar_trait_grid) == 12 and all(value == 15 for value in cultivar_trait_grid.values())
    audit.check("all 12 supplemental trait facets cover 15 cultivars", all_traits_have_15_cultivars,
                cultivar_trait_grid, {trait: 15 for trait in sorted(cultivar_trait_grid)})

    training = pd.read_csv(FIGURE_DATA / "training_dynamics.csv")
    training_histories = training[["trait", "outer_fold"]].drop_duplicates()
    training_ok = (
        len(training_histories) == 60
        and training["trait"].nunique() == 12
        and training["outer_fold"].nunique() == 5
        and training[["train_loss", "validation_residual_score", "learning_rate"]].notna().all().all()
    )
    audit.check("authentic training dynamics cover 12 traits x five folds", training_ok,
                {"histories": len(training_histories), "traits": training["trait"].nunique(),
                 "folds": training["outer_fold"].nunique()},
                {"histories": 60, "traits": 12, "folds": 5})
    auc_boundary = "ROC/AUC curves are not applicable" in supp_text and "continuous" in supp_text
    audit.check("ROC/AUC omission is explicitly bounded to regression endpoints",
                auc_boundary, auc_boundary, True)

    fig3d = pd.read_csv(FIGURE_DATA / "fig03d_fivefold_se.csv")
    fig3d_ok = (
        len(fig3d) == 36
        and fig3d["trait"].nunique() == 12
        and fig3d["baseline_display"].nunique() == 3
        and fig3d["n"].eq(5).all()
        and fig3d[["mean", "se"]].notna().all().all()
        and fig3d["se"].ge(0).all()
    )
    audit.check("Descriptive five-fold SE audit table is complete", fig3d_ok,
                {"rows": len(fig3d), "traits": fig3d["trait"].nunique(),
                 "comparators": fig3d["baseline_display"].nunique(),
                 "n_values": sorted(fig3d["n"].unique().tolist())},
                {"rows": 36, "traits": 12, "comparators": 3, "n_values": [5]})
    interval_widths = comparisons[[
        "trait", "baseline_display", "relative_improvement_ci_low", "relative_improvement_ci_high"
    ]].copy()
    interval_widths["bootstrap_half_width"] = (
        interval_widths["relative_improvement_ci_high"]
        - interval_widths["relative_improvement_ci_low"]
    ) / 2
    width_comparison = fig3d.merge(interval_widths, on=["trait", "baseline_display"])
    median_ratio = float((width_comparison["se"] / width_comparison["bootstrap_half_width"]).median())
    audit.check("Archived five-fold SE is materially smaller than full bootstrap uncertainty",
                median_ratio < 0.25, round(median_ratio, 6), "< 0.25")
    audit.require_file(SUPPLEMENT / "figures" / "figS17_training_dynamics.png",
                       "Supplementary Figure S17 training dynamics")
    audit.require_file(SUPPLEMENT / "figures" / "figS18_cluster_bootstrap_contrasts.png",
                       "Supplementary Figure S18 full bootstrap intervals")
    audit.require_file(SUPPLEMENT / "figures" / "figS19_FW_multiseed_robustness.png",
                       "Supplementary Figure S19 FW multiseed robustness")
    for index, label in [
        (20, "held-batch few-shot performance"),
        (21, "held-batch few-shot gain"),
        (22, "cultivar texture profiles"),
        (23, "multiplicity-adjusted contrasts"),
        (24, "cultivar QC decision"),
        (25, "all-trait multiseed robustness"),
    ]:
        candidates = sorted((SUPPLEMENT / "figures").glob(f"figS{index:02d}_*.png"))
        audit.check(f"Supplementary Figure S{index:02d} {label}", len(candidates) == 1,
                    [path.name for path in candidates], "exactly one PNG")
    v23_summary_path = ROOT / "results" / "v23_multiseed" / "analysis" / "fw_multiseed_summary.json"
    audit.require_file(v23_summary_path, "V23 FW multiseed summary")
    if v23_summary_path.exists():
        v23_summary = json.loads(v23_summary_path.read_text(encoding="utf-8"))
        v23_ok = (
            v23_summary.get("total_fits") == 25
            and v23_summary.get("seeds_per_fold") == 5
            and v23_summary.get("outer_folds") == 5
            and v23_summary.get("unique_test_fruits") == 4828
            and v23_summary.get("multiseed_fold_wins_vs_domain_pls") == 3
            and v23_summary.get("se_reduction_pct", 0) > 40
        )
        audit.check("V23 FW multiseed audit has complete non-pseudoreplicated coverage", v23_ok,
                    v23_summary, {"fits": 25, "folds": 5, "seeds_per_fold": 5,
                                  "fruits": 4828, "SE reduction": ">40%", "fold wins": 3})

    v24_summary_path = V24 / "analysis" / "v24_hr_strengthening_summary.json"
    alltrait_summary_path = V24 / "multiseed_analysis" / "alltrait_multiseed_summary.json"
    all12_path = V24 / "multiseed_analysis" / "all12_multiseed_summary.csv"
    for path, label in [
        (v24_summary_path, "V24 strengthening summary"),
        (alltrait_summary_path, "V24 all-trait multiseed summary"),
        (all12_path, "V24 all-12 multiseed summary"),
    ]:
        audit.require_file(path, label)
    if v24_summary_path.exists():
        v24_summary = json.loads(v24_summary_path.read_text(encoding="utf-8"))
        v24_ok = (
            v24_summary["fewshot"]["fruits"] == 1236
            and v24_summary["fewshot"]["traits"] == 9
            and v24_summary["fewshot"]["repeats"] == 500
            and v24_summary["multiplicity"]["contrasts"] == 36
            and v24_summary["multiplicity"]["supported_simultaneous"] == 24
            and v24_summary["cultivar_texture"]["cultivars"] == 15
            and v24_summary["cultivar_texture"]["traits"] == 9
            and v24_summary["quality_control"]["residuals_used_for_exclusion"] is False
        )
        audit.check("V24 strengthening evidence is complete", v24_ok, v24_summary,
                    "1236 fruits; 9 held-batch traits; 500 resamples; 36 contrasts; 24 simultaneous; 15 cultivars; no residual QC")
    if alltrait_summary_path.exists() and all12_path.exists():
        alltrait_summary = json.loads(alltrait_summary_path.read_text(encoding="utf-8"))
        all12 = pd.read_csv(all12_path)
        non_fw = all12.loc[all12["trait"] != "FW"].copy()
        multiseed_ok = (
            alltrait_summary.get("additional_fits") == 110
            and alltrait_summary.get("pipeline_instances_including_original") == 165
            and len(all12) == 12
            and int(non_fw["cluster_supported"].sum()) == 11
            and int(all12.loc[all12["trait"] == "FW", "cluster_supported"].iloc[0]) == 0
            and float(non_fw["mean_within_fold_seed_sd_pct_points"].median()) < 0.25
        )
        audit.check("V24 multiseed robustness supports 11/12 targets with FW at parity", multiseed_ok,
                    {"additional_fits_non_FW": alltrait_summary.get("additional_fits"),
                     "all12_rows": len(all12), "supported_non_FW": int(non_fw["cluster_supported"].sum()),
                     "median_seed_SD_pp": round(float(non_fw["mean_within_fold_seed_sd_pct_points"].median()), 6)},
                    {"additional_fits_non_FW": 110, "all12_rows": 12, "supported_non_FW": 11,
                     "median_seed_SD_pp": "<0.25"})

    for path, label in [
        (ROOT / "review_package" / "16_V24_HR_STRENGTHENING_REPORT_ZH.md", "V24 HR Chinese project report"),
        (MANUSCRIPT / "Horticulture_Research_submission_information_form.docx", "author information form"),
        (MANUSCRIPT / "Horticulture_Research_cover_letter_draft.docx", "HR cover-letter draft"),
    ]:
        audit.require_file(path, label)

    package_counts = {
        "main_figures": len(list((HR_PACKAGE / "main_figures").glob("*"))),
        "supplementary_vector_figures": len(list((HR_PACKAGE / "supplementary_vector_figures").glob("*.pdf"))),
        "supplementary_tables": len(list((HR_PACKAGE / "supplementary_tables").glob("*.csv"))),
    }
    audit.check("HR hand-off package payload counts",
                package_counts == {"main_figures": 7, "supplementary_vector_figures": 25,
                                   "supplementary_tables": 32},
                package_counts,
                {"main_figures": 7, "supplementary_vector_figures": 25, "supplementary_tables": 32})
    checksum_path = HR_PACKAGE / "SHA256SUMS.csv"
    audit.require_file(checksum_path, "HR hand-off package SHA-256 manifest")
    if checksum_path.exists():
        checksum_rows = pd.read_csv(checksum_path)
        checksum_ok = len(checksum_rows) == 71 and checksum_rows["sha256"].str.fullmatch(r"[0-9a-f]{64}").all()
        audit.check("HR hand-off package checksum coverage", checksum_ok,
                    {"rows": len(checksum_rows), "valid_sha256": bool(checksum_rows["sha256"].str.fullmatch(r"[0-9a-f]{64}").all())},
                    {"rows": 71, "valid_sha256": True})

    result = audit.finish()
    AUDIT_DIR.mkdir(parents=True, exist_ok=True)
    AUDIT_JSON.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result["summary"], ensure_ascii=False))
    print(f"status={result['status']}")
    print(f"report={AUDIT_JSON}")
    if result["status"] != "PASS":
        for item in result["checks"]:
            if item["status"] == "FAIL":
                print(f"FAIL {item['name']}: observed={item['observed']} expected={item['expected']}")
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
