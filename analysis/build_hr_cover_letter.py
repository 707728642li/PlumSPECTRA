from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "manuscript" / "Horticulture_Research_cover_letter_draft.docx"
NAVY = RGBColor(23, 36, 59)
PLUM = RGBColor(123, 45, 95)
MUTED = RGBColor(92, 105, 124)
YELLOW = "FFF2CC"


def set_font(run, size: float = 10.5, *, bold: bool = False,
             italic: bool = False, colour: RGBColor = NAVY) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = colour


def highlight(run) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), YELLOW)
    run._r.get_or_add_rPr().append(shading)


def add_paragraph(doc: Document, text: str, *, space_after: float = 7,
                  bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_font(rest)
    else:
        run = paragraph.add_run(text)
        set_font(run)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = NAVY

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(3)
    run = heading.add_run("COVER LETTER")
    set_font(run, 15, bold=True, colour=PLUM)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(12)
    run = sub.add_run("Horticulture Research — Article submission draft")
    set_font(run, 9.5, bold=True, colour=MUTED)

    date = doc.add_paragraph()
    date.paragraph_format.space_after = Pt(7)
    run = date.add_run("[SUBMISSION DATE]")
    set_font(run)
    highlight(run)

    add_paragraph(doc, "Dear Editors of Horticulture Research,")

    title = (
        'We submit the Article “PlumSPECTRA unifies nondestructive prediction of '
        'conventional quality and nine mechanical texture phenotypes across 4,839 plums” '
        'for consideration in Horticulture Research.'
    )
    add_paragraph(doc, title)

    add_paragraph(
        doc,
        "The study moves fruit near-infrared analysis beyond soluble solids and a single firmness value. "
        "We linked intact-fruit spectra to duplicate texture-analyser curves, fruit weight, soluble solids "
        "and pH across 15 retained plum cultivars. Nine mechanically interpretable endpoints describe peel "
        "rupture, deformation, flesh resistance, stiffness, penetration work and withdrawal adhesion. The "
        "resulting phenotype atlas reveals cultivar-associated mechanical diversity that a single firmness "
        "measurement would conceal.",
    )

    add_paragraph(
        doc,
        "PlumSPECTRA fits one compact residual convolutional model per trait around a cultivar-aware "
        "chemometric anchor and, when selected using training data only, a nonlinear kernel branch. Frozen "
        "cultivar-stratified outer folds produced 58,035 out-of-fold fruit–trait predictions. The final system "
        "reduced error for all 12 targets relative to global partial least squares, cultivar-aware partial "
        "least squares and nested support-vector regression; multiplicity-adjusted cultivar-cluster inference "
        "separates the strongest conclusions from smaller gains. Duplicate texture reliability, within-cultivar "
        "metrics, complete-pipeline seed repeats and an explicit whole-cultivar quality audit are reported.",
    )

    add_paragraph(
        doc,
        "We also disclose a stringent deployment boundary. Direct prediction failed when complete same-cultivar "
        "acquisition batches were withheld. Reference-based recalibration with 40 labelled fruit per batch "
        "restored positive pooled performance for most texture traits, while batch-level heterogeneity remained. "
        "This result converts a common but often hidden limitation of fruit spectroscopy into a testable "
        "calibration strategy and a prospective design for new-season validation.",
    )

    add_paragraph(
        doc,
        "The manuscript fits the journal because it combines large-scale phenotyping of a major horticultural "
        "crop with cultivar-resolved mechanical diversity, breeding/postharvest relevance and a rigorously "
        "validated nondestructive platform. The central contribution is the multidimensional plum phenotype "
        "resource and its scalable spectral surrogate, rather than architecture alone.",
    )

    add_paragraph(
        doc,
        "We confirm that the work is original, is not under consideration elsewhere and has been approved by "
        "all authors. Data, code, frozen split manifests and reviewer-access details will be available at "
        "[REPOSITORY DOI OR REVIEWER LINK].",
    )
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "[REPOSITORY DOI OR REVIEWER LINK]" in run.text:
                highlight(run)

    add_paragraph(doc, "Thank you for considering our work.")
    add_paragraph(doc, "Sincerely,", space_after=3)
    signature = doc.add_paragraph()
    signature.paragraph_format.space_after = Pt(0)
    for line in ("[CORRESPONDING AUTHOR]", "[AFFILIATION]", "[EMAIL]", "[POSTAL ADDRESS]"):
        run = signature.add_run(line + "\n")
        set_font(run)
        highlight(run)

    doc.core_properties.title = "Horticulture Research cover letter draft"
    doc.core_properties.author = "Anonymous draft"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
