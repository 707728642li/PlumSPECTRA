from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(__file__).resolve().parents[1]
FIGURES = {
    1: PROJECT / "results/v20/figures_hr/fig01_v20_study_design.png",
    2: PROJECT / "results/v20/figures_hr/fig02_texture_phenotype_atlas.png",
    3: PROJECT / "results/v20/figures_hr/fig03_v20_model_performance.png",
    4: PROJECT / "results/v20/figures_hr/fig04_v20_all_trait_predictions.png",
    5: PROJECT / "results/v20/figures_hr/fig05_v20_within_cultivar_heterogeneity.png",
    6: PROJECT / "results/v20/figures_hr/fig06_v21_crossbatch_boundary.png",
}

INK = "20262E"
BLUE = "2E5C87"
DARK_BLUE = "1F405F"
MUTED = "65717D"


def set_run_font(run, size: float, *, bold: bool = False, italic: bool = False, color: str = INK) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_style(style, *, size: float, before: float, after: float, color: str, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.keep_with_next = True


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_run_font(run, 9, color=MUTED)


def set_document_tokens(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    configure_style(doc.styles["Heading 1"], size=16, before=18, after=14, color=BLUE, bold=True)
    configure_style(doc.styles["Heading 2"], size=13, before=12, after=10, color=BLUE, bold=True)
    configure_style(doc.styles["Heading 3"], size=12, before=8, after=4, color=DARK_BLUE, bold=True)


def add_numbering_definition(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    # Named reference-list override: two-digit citation markers need more room
    # than the narrative preset's single-level body list.
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    level.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


INLINE = re.compile(r"(\*\*.*?\*\*|\*[^*]+?\*|`.*?`)")


def add_inline(paragraph, text: str, *, size: float = 11, color: str = INK) -> None:
    cursor = 0
    for match in INLINE.finditer(text):
        if match.start() > cursor:
            set_run_font(paragraph.add_run(text[cursor : match.start()]), size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size, bold=True, color=color)
        elif token.startswith("*"):
            set_run_font(paragraph.add_run(token[1:-1]), size, italic=True, color=color)
        else:
            set_run_font(paragraph.add_run(token[1:-1]), size - 0.5, color=DARK_BLUE)
        cursor = match.end()
    if cursor < len(text):
        set_run_font(paragraph.add_run(text[cursor:]), size, color=color)


def add_cover(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(92)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("HORTICULTURE RESEARCH | REVIEW MANUSCRIPT"), 10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(title), 25, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Anonymous scientific review draft"), 12, italic=True, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(130)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Frozen V20 five-fold audit + sealed V21 cross-batch boundary test"), 10.5, color=MUTED)
    doc.add_page_break()


def add_figure(doc: Document, number: int) -> None:
    path = FIGURES[number]
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    # Do not chain the full-resolution figure to the following legend.  Word's
    # pagination engine can otherwise pull a heading/picture/legend chain past
    # the page boundary and clip the first heading line after a manual break.
    p.paragraph_format.keep_with_next = False
    shape = p.add_run().add_picture(str(path), width=Inches(6.25))
    shape._inline.docPr.set("title", f"Figure {number}")
    shape._inline.docPr.set("descr", f"Scientific figure {number}; full alternative text follows the legend.")


def build(input_path: Path, output_path: Path) -> None:
    lines = input_path.read_text(encoding="utf-8").splitlines()
    title = next(line[2:].strip() for line in lines if line.startswith("# "))
    doc = Document()
    set_document_tokens(doc)
    num_id = add_numbering_definition(doc)

    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run("Plum NIR texture phenotyping | review draft"), 8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(footer)
    add_cover(doc, title)

    in_references = False
    in_figure_legends = False
    paragraph_buffer: list[str] = []

    def flush() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer).strip()
        paragraph_buffer = []
        if not text:
            return
        reference = re.match(r"^\d+\.\s+(.*)$", text) if in_references else None
        p = doc.add_paragraph()
        if reference:
            apply_numbering(p, num_id)
            add_inline(p, reference.group(1), size=9.5)
        else:
            add_inline(p, text)

    for line in lines:
        hard_break = line.endswith("  ")
        stripped = line.strip()
        if stripped.startswith("# "):
            continue
        if stripped.startswith("## "):
            flush()
            heading = stripped[3:].strip()
            in_references = heading == "References"
            in_figure_legends = heading == "Figure legends"
            if in_figure_legends:
                doc.add_page_break()
            heading_p = doc.add_paragraph(heading, style="Heading 1")
            # Keep the abstract as a self-contained front-matter page and avoid
            # crowding the Introduction into its last line.
            if heading == "Introduction":
                heading_p.paragraph_format.page_break_before = True
            continue
        if stripped.startswith("### "):
            flush()
            heading = stripped[4:].strip()
            if in_figure_legends:
                match = re.match(r"Figure\s+(\d+)\.", heading)
                if match:
                    number = int(match.group(1))
                    label, title = heading.split(". ", 1)
                    # Figure sheets are already page-isolated.  Direct paragraph
                    # formatting avoids a Word pagination quirk in which a
                    # keep-with-next Heading 2 can be placed above the printable
                    # area on image-dense pages (observed for Figures 3 and 5).
                    label_p = doc.add_paragraph()
                    # Figure 3 is the densest sheet. Word's PDF exporter can
                    # place its first baseline one text-line above the page box;
                    # reserve an explicit top offset so the review PDF and DOCX
                    # remain visually identical and unclipped.
                    label_p.paragraph_format.space_before = Pt(84 if number == 3 else 12)
                    label_p.paragraph_format.space_after = Pt(0)
                    label_p.paragraph_format.keep_with_next = False
                    label_p.paragraph_format.left_indent = Inches(0)
                    label_p.paragraph_format.first_line_indent = Inches(0)
                    set_run_font(label_p.add_run(f"{label}."), 13, bold=True, color=BLUE)
                    if number > 1:
                        # A paragraph-level page break is more robust than a
                        # separate empty break paragraph in Word/LibreOffice.
                        label_p.paragraph_format.page_break_before = True
                    title_p = doc.add_paragraph()
                    title_p.paragraph_format.space_before = Pt(0)
                    title_p.paragraph_format.space_after = Pt(6)
                    title_p.paragraph_format.keep_with_next = False
                    title_p.paragraph_format.left_indent = Inches(0)
                    title_p.paragraph_format.first_line_indent = Inches(0)
                    set_run_font(title_p.add_run(title), 13, bold=True, color=BLUE)
                    add_figure(doc, number)
                    continue
            doc.add_paragraph(heading, style="Heading 2")
            continue
        if not stripped:
            flush()
            continue
        if in_references and re.match(r"^\d+\.\s+", stripped):
            flush()
            paragraph_buffer.append(stripped)
            flush()
        elif hard_break:
            paragraph_buffer.append(stripped)
            flush()
        else:
            paragraph_buffer.append(stripped)
    flush()

    # Re-apply geometry to every section created by Word while preserving the
    # exact narrative-proposal preset tokens.
    set_document_tokens(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = title
    doc.core_properties.subject = "Nondestructive NIR prediction of nine plum texture phenotypes"
    doc.core_properties.author = "Anonymous review draft"
    doc.core_properties.keywords = "plum; NIR spectroscopy; texture; deep learning; PLSR; SVR"
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build(args.input.resolve(), args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
