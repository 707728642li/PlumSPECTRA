from __future__ import annotations

import csv
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
OUTPUT = MANUSCRIPT / "Cultivar_shift_plum_NIR_Horticulture_Research.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5F6B78"
LIGHT_FILL = "F4F6F9"
RULE = "A8B2BD"


FIGURES = {
    "fig1": (
        ROOT / "results/figures_main/fig01_study_design.png",
        "Figure 1. Cohort construction and leakage-safe evaluation design. "
        "(a) Flow from source inventory through structural curation, target-specific NIR cohorts and the strict complete multimodal core. "
        "(b) Individual-fruit counts for all 16 cultivar or breeding-selection labels. "
        "(c) Roles of NIR, reference traits and duplicate texture curves in leave-one-cultivar-out benchmarking, few-shot adaptation and conformal intervals.",
    ),
    "fig2": (
        ROOT / "results/eda/figures/fig02_phenotypic_spectral_diversity.png",
        "Figure 2. Cultivar structure in plum phenotypes and NIR spectra. Distributions of weight, SSC and pH, representative spectra and principal-component separation show that cultivar affects both response and predictor spaces. Percent variance attributable to cultivar is η² from one-way sums of squares.",
    ),
    "fig3": (
        ROOT / "results/eda/figures/fig03_texture_reliability_biology.png",
        "Figure 3. Automated ARC curve phenotyping reveals repeatable, cultivar-structured mechanics. "
        "(a) Absolute-agreement single-measure ICC across duplicate penetrations. "
        "(b) Cultivar median z-scored texture fingerprints. "
        "(c) Global Spearman associations of texture features with fruit weight, pH and SSC. "
        "(d) Within-cultivar Spearman associations after removing between-cultivar mean structure.",
    ),
    "fig4": (
        ROOT / "results/model_comparison/figures/fig04_zero_shot_model_comparison.png",
        "Figure 4. Zero-shot prediction deteriorates when the test cultivar is absent from training. Pooled and cultivar-level performance of nested PLSR, three-seed CNN, Transformer and texture-auxiliary CNN under identical leave-one-cultivar-out folds. Random-split PLSR is shown only as an interpolation benchmark.",
    ),
    "fig5": (
        ROOT / "results/model_comparison/figures/fig05_fewshot_calibration_curves.png",
        "Figure 5. Few-shot intercept calibration restores prediction in a new cultivar. Repeated adaptation curves for direct PLSR, deep models and hierarchical PLS. Lines and bands show means and 95% empirical intervals across 100 identical calibration draws; calibration fruits are excluded from evaluation.",
    ),
    "fig6": (
        ROOT / "results/model_comparison/figures/fig06_hierarchical_10shot_predictions.png",
        "Figure 6. Representative ten-shot hierarchical predictions. Observed versus predicted values after intercept calibration with ten labelled fruits per held-out cultivar in repeat 1. The identity line indicates perfect agreement; displayed metrics apply only to evaluation fruits.",
    ),
}

ALT_TEXT = {
    "fig1": "A three-part study diagram. A left-side flow chart reduces 5,894 source records to 5,487 complete fruits, a right-side horizontal bar chart lists sample counts for 16 cultivar or selection labels, and a lower workflow separates zero-shot benchmarking, few-shot adaptation and prediction intervals.",
    "fig2": "A multi-panel diversity plot showing cultivar distributions of fruit weight, soluble solids and pH, overlaid near-infrared spectral profiles and a principal-component score plot. Cultivars occupy different phenotype ranges and partially separated spectral regions.",
    "fig3": "A four-panel texture plot showing bars for duplicate-measure reliability, a heat map of cultivar median mechanical fingerprints, a heat map of global mechanical-quality correlations and a heat map of within-cultivar correlations after removing between-cultivar mean structure.",
    "fig4": "Three target-specific panels compare zero-shot errors for direct partial least-squares regression, hierarchical partial least-squares regression, convolutional networks and a Transformer. Performance varies widely among held-out cultivars and deep models do not consistently improve on direct partial least-squares regression.",
    "fig5": "Three line charts show prediction performance as the number of labelled target-cultivar fruits increases from zero to 50. Hierarchical calibration rises sharply by five to ten fruits and then plateaus, especially for fruit weight.",
    "fig6": "Three observed-versus-predicted scatter plots for fruit weight, soluble solids and pH after ten-fruit hierarchical intercept calibration. Points cluster around identity more closely than in zero-shot evaluation, with remaining cultivar-level spread.",
}


TABLES = {
    "table1": (
        MANUSCRIPT / "tables/table1_cohort.csv",
        "Table 1. Strict complete multimodal cohort by cultivar or breeding selection.",
        [1750, 650, 650, 2200, 2050, 2060],
        "IQR, interquartile range; SSC, soluble solids concentration. Counts require valid weight, SSC, pH, a valid primary NIR spectrum and two valid texture curves.",
    ),
    "table2": (
        MANUSCRIPT / "tables/table2_model_performance.csv",
        "Table 2. Random-split interpolation and zero-shot leave-one-cultivar-out performance.",
        [1850, 1250, 650, 1122, 1122, 1122, 1122, 1122],
        "Deep-learning values are ensembles across three fixed seeds. Random-split PLSR values are means across five repeated cultivar-stratified 80:20 splits; unavailable aggregate fields are shown as em dashes.",
    ),
    "table3": (
        MANUSCRIPT / "tables/table3_fewshot_hierarchical.csv",
        "Table 3. Few-shot intercept calibration of hierarchical PLSR in held-out cultivars.",
        [900, 1400, 3900, 1300, 1860],
        "Values summarize 100 repeated calibration draws; calibration fruits were excluded from evaluation. At zero shots the interval collapses because the same zero-shot predictions were repeated for comparability.",
    ),
    "table4": (
        MANUSCRIPT / "tables/table4_conformal_intervals.csv",
        "Table 4. Within-held-out-cultivar split-conformal intervals at nominal 90% coverage.",
        [1050, 950, 950, 1150, 1650, 1900, 1710],
        "Half of labelled fruits fitted the intercept and half fitted the absolute-residual quantile. All labelled fruits were excluded from evaluation. Width units follow the target (g, % SSC or pH units).",
    ),
}


REFERENCE_DOIS = [
    "10.1016/j.postharvbio.2007.06.024",
    "10.1016/j.postharvbio.2020.111139",
    "10.5935/0103-5053.20130172",
    "10.1111/jfpe.13597",
    None,
    "10.1016/j.saa.2026.128279",
    "10.1016/j.saa.2023.123151",
    "10.1111/jfpp.16504",
    "10.1016/j.postharvbio.2020.111202",
    "10.1016/j.chemolab.2021.104287",
    "10.1016/j.saa.2024.124003",
    "10.1016/j.chemolab.2023.104924",
    "10.1016/j.postharvbio.2024.112783",
    "10.1016/j.saa.2025.126122",
    "10.1016/j.compag.2026.112186",
    "10.1016/j.aiia.2025.12.003",
    "10.1016/j.foodcont.2024.110823",
    "10.3389/fpls.2023.1128993",
    "10.1093/fqsafe/fyac068",
    "10.1016/j.foodchem.2025.145387",
    "10.1016/j.scienta.2014.01.002",
    "10.1016/j.scienta.2026.114793",
    "10.1038/s41438-021-00560-9",
    "10.1016/j.foodchem.2026.148344",
    "10.1561/2200000101",
    "10.1021/ac60214a047",
    "10.1016/j.jcm.2016.02.012",
    "10.1109/CVPR.2016.90",
    None,
    "10.2307/2532051",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != 9360:
        raise ValueError(f"Table widths must sum to 9360 DXA, got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "C8CDD3")

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_run(run, size: float = 11, bold: bool = False, italic: bool = False, color: str = "000000") -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def add_inline(paragraph, text: str, size: float = 11) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            set_run(paragraph.add_run(text[position : match.start()]), size=size)
        token = match.group(0)
        if token.startswith("**"):
            set_run(paragraph.add_run(token[2:-2]), size=size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run(run, size=size - 0.5, color=DARK_BLUE)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        position = match.end()
    if position < len(text):
        set_run(paragraph.add_run(text[position:]), size=size)


def add_body(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    add_inline(paragraph, text)


def add_figure(doc: Document, key: str) -> None:
    path, caption = FIGURES[key]
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    shape = paragraph.add_run().add_picture(str(path), width=Inches(6.35))
    shape._inline.docPr.set("title", f"Figure {key.removeprefix('fig')}")
    shape._inline.docPr.set("descr", ALT_TEXT[key])
    cap = doc.add_paragraph(style="Caption")
    cap.paragraph_format.keep_with_next = True
    add_inline(cap, caption, size=9)
    alt = doc.add_paragraph(style="Table Note")
    add_inline(alt, f"Alt text: {ALT_TEXT[key]}", size=8.5)


def add_table(doc: Document, key: str) -> None:
    path, caption, widths, note = TABLES[key]
    frame = pd.read_csv(path, dtype=str).fillna("—")
    cap = doc.add_paragraph(style="Table Caption")
    add_inline(cap, caption, size=9.5)
    table = doc.add_table(rows=1, cols=len(frame.columns))
    set_repeat_table_header(table.rows[0])
    for index, value in enumerate(frame.columns):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(value))
        set_run(run, size=7.5, bold=True, color=DARK_BLUE)
    for record in frame.itertuples(index=False, name=None):
        cells = table.add_row().cells
        for index, value in enumerate(record):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index in (0, 1) else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            set_run(run, size=7.3)
    set_table_geometry(table, widths)
    note_p = doc.add_paragraph(style="Table Note")
    add_inline(note_p, f"Note: {note}", size=8.5)


def format_author(name: str) -> str:
    name = " ".join(name.split())
    if "," in name:
        family, given = [part.strip() for part in name.split(",", 1)]
    else:
        parts = name.split()
        particle_start = len(parts) - 1
        while particle_start > 0 and parts[particle_start - 1].lower() in {
            "da",
            "de",
            "del",
            "der",
            "di",
            "dos",
            "la",
            "van",
            "von",
        }:
            particle_start -= 1
        family = " ".join(parts[particle_start:])
        given = " ".join(parts[:particle_start])
    initial_parts: list[str] = []
    for token in given.split():
        letters = "".join(character for character in token if character.isalpha())
        if not letters:
            continue
        if token.count(".") >= 2 or (letters.isupper() and len(letters) <= 4):
            initial_parts.append(letters.upper())
        else:
            initial_parts.append(letters[0].upper())
    initials = "".join(initial_parts)
    return f"{family} {initials}".strip()


def reference_rows() -> list[str]:
    metadata = pd.read_csv(ROOT / "literature/reference_metadata.csv", dtype=str).fillna("")
    by_doi = {row.doi.lower(): row for row in metadata.itertuples(index=False)}
    rows: list[str] = []
    manual_index = 0
    for doi in REFERENCE_DOIS:
        if doi is None:
            if manual_index == 0:
                rows.append(
                    "DI B, LIN J, LIU X. Research on prediction model of soluble solid content in plums based on near-infrared spectroscopy data. Chinese Agricultural Science Bulletin. 2024;40(31):133–138. doi:10.11924/j.issn.1000-6850.casb2024-0452."
                )
            else:
                rows.append(
                    "Vaswani A, Shazeer N, Parmar N, et al. Attention is all you need. Advances in Neural Information Processing Systems. 2017;30:5998–6008. arXiv:1706.03762."
                )
            manual_index += 1
            continue
        record = by_doi[doi.lower()]
        authors = [format_author(name.strip()) for name in record.authors.split(";") if name.strip()]
        author_text = ", ".join(authors[:6]) + (", et al." if len(authors) > 6 else ".")
        volume = str(record.volume)
        issue = f"({record.issue})" if str(record.issue) else ""
        pages = f":{record.pages_or_article}" if str(record.pages_or_article) else ""
        title = str(record.title).strip()
        title_text = title if title.endswith((".", "?", "!")) else f"{title}."
        journal = str(record.journal).strip()
        journal_text = journal if journal.endswith((".", "?", "!")) else f"{journal}."
        rows.append(
            f"{author_text} {title_text} {journal_text} {record.year};{volume}{issue}{pages}. doi:{record.doi}."
        )
    return rows


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string("30363D")
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.line_spacing = 1.05

    for name in ("Table Caption", "Table Note"):
        if name not in styles:
            styles.add_style(name, 1)
    table_caption = styles["Table Caption"]
    table_caption.font.name = "Calibri"
    table_caption.font.size = Pt(9.5)
    table_caption.font.bold = True
    table_caption.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    table_caption.paragraph_format.space_before = Pt(8)
    table_caption.paragraph_format.space_after = Pt(4)
    table_caption.paragraph_format.keep_with_next = True
    table_note = styles["Table Note"]
    table_note.font.name = "Calibri"
    table_note.font.size = Pt(8.5)
    table_note.font.italic = True
    table_note.font.color.rgb = RGBColor.from_string(MUTED)
    table_note.paragraph_format.space_before = Pt(4)
    table_note.paragraph_format.space_after = Pt(8)


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("FEW-SHOT TRANSFER OF PLUM NIR MODELS")
    set_run(run, size=8.5, color=MUTED)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), RULE)
    borders.append(bottom)
    p_pr.append(borders)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer_p.add_run("Page "), size=8.5, color=MUTED)
    page_run = footer_p.add_run()
    set_run(page_run, size=8.5, color=MUTED)
    add_field(page_run, "PAGE")


def add_title_page(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("HORTICULTURE RESEARCH · RESEARCH ARTICLE"), size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    set_run(p.add_run(title), size=24, bold=True, color=DARK_BLUE)

    for text, size, italic in [
        ("[AUTHOR NAMES]", 12, False),
        ("[AFFILIATIONS]", 10.5, True),
        ("Author emails: [EMAIL ADDRESS FOR EACH AUTHOR]", 10, False),
        ("Corresponding author: [NAME, POSTAL ADDRESS AND EMAIL]", 10, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        set_run(p.add_run(text), size=size, italic=italic, color=MUTED if italic else "000000")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run("Submission metadata to complete"), size=10.5, bold=True, color=DARK_BLUE)
    items = [
        "Authors, affiliations, corresponding-author details and CRediT roles",
        "NIR instrument model/optical geometry and orchard, season and maturity metadata",
        "SSC and pH reference instruments and laboratory procedures",
        "Funding, acknowledgements, competing interests and public repository DOIs",
    ]
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(item), size=9.5, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Prepared from the independently audited research-ready release v1.1.0"), size=9, italic=True, color=MUTED)
    doc.add_page_break()


def main() -> None:
    source_path = MANUSCRIPT / "manuscript_final.md"
    if not source_path.exists():
        raise FileNotFoundError("Run build_manuscript_assets.py first")
    source = source_path.read_text(encoding="utf-8")
    title = source.splitlines()[0].removeprefix("# ").strip()

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_title_page(doc, title)

    inserted: set[str] = set()
    started = False
    skip_legends = False
    for raw in source.splitlines():
        line = raw.strip()
        if line == "## Abstract":
            started = True
        if not started:
            continue
        if line == "## Figure legends":
            skip_legends = True
            continue
        if skip_legends:
            continue
        if not line:
            continue
        if line == "## References":
            doc.add_heading("References", level=1)
            for index, reference in enumerate(reference_rows(), start=1):
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.28)
                paragraph.paragraph_format.first_line_indent = Inches(-0.28)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.05
                set_run(paragraph.add_run(f"{index}. "), size=8.8, bold=True, color=DARK_BLUE)
                set_run(paragraph.add_run(reference), size=8.8)
            continue
        if line.startswith("References are generated"):
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            continue
        if line.startswith("# "):
            continue

        add_body(doc, line)

        if "removing them did not materially change any headline RMSE" in line and "cohort" not in inserted:
            add_figure(doc, "fig1")
            add_table(doc, "table1")
            inserted.add("cohort")
        elif "This joint structure predicts optimistic validation" in line and "spectral" not in inserted:
            add_figure(doc, "fig2")
            inserted.add("spectral")
        elif "optical transfer occurs against a background of strongly cultivar-dependent tissue mechanics" in line and "texture" not in inserted:
            add_figure(doc, "fig3")
            inserted.add("texture")
        elif "this experiment estimates interpolation" in line and "zero_table" not in inserted:
            add_table(doc, "table2")
            inserted.add("zero_table")
        elif "demonstrates a universal plum model" in line and "zero_fig" not in inserted:
            add_figure(doc, "fig4")
            inserted.add("zero_fig")
        elif "principal calibration sizes" in line and "fewshot_table" not in inserted:
            add_figure(doc, "fig5")
            add_table(doc, "table3")
            inserted.add("fewshot_table")
        elif "same reference methods used to define deployment targets" in line and "fewshot_fig6" not in inserted:
            add_figure(doc, "fig6")
            inserted.add("fewshot_fig6")
        elif "evaluated interval procedure" in line and "conformal" not in inserted:
            add_table(doc, "table4")
            inserted.add("conformal")

    missing = {
        "cohort",
        "spectral",
        "texture",
        "zero_table",
        "zero_fig",
        "fewshot_table",
        "fewshot_fig6",
        "conformal",
    } - inserted
    if missing:
        raise RuntimeError(f"Failed to insert manuscript assets: {sorted(missing)}")
    doc.core_properties.title = title
    doc.core_properties.subject = "Cultivar-held-out NIR phenotyping of plum quality"
    doc.core_properties.keywords = "plum; NIR; cultivar shift; few-shot calibration; texture; conformal prediction"
    doc.core_properties.author = "[AUTHOR NAMES]"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
